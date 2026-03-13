"""
TODO: support trade studies across single or multiple axes
TODO: x-axis can be mach or kact
TODO: y-axis can be eas or frequency
"""
from __future__ import annotations
import os
import warnings
from pathlib import Path
from itertools import count
from typing import Callable, Optional, TYPE_CHECKING

import numpy as np
from matplotlib import pyplot as plt
# import matplotlib.gridspec as gridspec

from pyNastran.utils import PathLike
from pyNastran.f06.flutter_response import get_damping_crossings as _get_damping_crossings, Limit
from pyNastran.f06.parse_flutter import make_flutter_response, FlutterResponse
from pyNastran.f06.dev.flutter.utils import get_vlines, get_noline_nopoints
from pyNastran.f06.dev.flutter.utils_report import write_docx_path

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    warnings.warn('>>> pip install python-docx')
    raise

if TYPE_CHECKING:
    import pandas as pd
    from cpylog import SimpleLogger


def write_report(docx_filename: str,
                 f06_filenames: list[str],
                 configs: list[str],
                 table: pd.DataFrame,
                 trades: list[dict],
                 log: SimpleLogger,
                 settings: dict[str, int | float | str],
                 x_plot_type: str='eas',
                 f06_units: str='english_in',
                 out_units: str='english_kt',
                 nrigid_body_modes: int=0,
                 modes=None,
                 vl_target: float=-1.0,
                 vf_target: float=-1.0,
                 # v_lines=None,
                 # xlim_kfreq=None,
                 ylim_damping: Limit=None,
                 ylim_freq: Limit=None,
                 eas_lim: Limit=None,
                 freq_tol: float=-1.0,
                 freq_tol_remove: float=-1.0,
                 damping_required: float=-1.0,        # VL
                 damping_required_tol: float=0.0001,  # VL
                 damping_limit: float=-1.0,           # VF
                 eas_flutter_range: Limit=None,
                 plot_font_size: int=10,
                 show_lines: bool=True,
                 show_points: bool=True,
                 show_mode_number: bool=False,
                 show_detailed_mode_info: bool=False,
                 point_spacing: int=8,
                 use_rhoref: bool=False,
                 flutter_ncolumns: int=0,
                 flutter_bbox_to_anchor_x=None,
                 freq_ndigits: int=2,
                 divergence_legend_loc: str='best',
                 divergence_freq_tol: float=0.1,
                 ndir_levels: int=1,
                 eas_max: float=1000.0,
                 progress_callback: Optional[Callable]=None,
                 ) -> None:
    """
    Parameters
    ----------
    docx_filename : str
        path to output file
    f06_filenames : str
        list of f06 files to process
    configs: list[str]
        label for the f06 file
    table: pd.DataFrame
        information for trade studies
    log : SimpleLogger
        logging object
    settings : dict[str, Any]
        the following parameters as a dictionary
    x_plot_type : str
        type of plot to make
    f06_units : str
        input units
    out_units : str
        output units
    nrigid_body_modes : int; default=0
        used with x_cutoff
    modes : np.ndarray; default=None -> all modes
        modes to consider
    vl_target : float; default=-1.0
        required speed for VL
        adds a solid black line
        adds a dotted black 1.15*VL line
        -1.0 is disabled
    vf_target : float; default=-1.0
        user red line to indicate flutter
        typically empty
        -1.0 is disabled
    ylim_damping : Limit
        damping range
    ylim_freq : Limit
        freqeuncy range
    eas_lim : Limit
        x range in equivalent airspeed
    freq_tol : float
        dashed line tolerance
    freq_tol_remove : float
        remove line tolerance
    damping_required : float; default=None
        the required damping at VL
        for a 0% crossing; damping_required=0.0
        for a -3% crossing; damping_required=-0.03
    damping_required_tol : float; default=0.0
        tolerance for damping_required
    damping_limit : float; default=None
        the required damping at VF
        for a 3% crossing; damping_limit=0.03
        for a 0% crossing; damping_limit=0.0
    eas_flutter_range : Limit
        the range of speeds to consider for flutter
        useful for getting rid of severe mode switching
        or rigid body crossings
    plot_font_size : int; default=8
        the size of the font
    show_lines : bool; default=True
        show the lines
    show_points : bool; default=True
        show the data points
    show_mode_number : bool; default=False
        show mode number instead of point
    show_detailed_mode_info : bool; default=False
        annotates hump modes with frequency & velocity range,
        as well as peak velocity/frequency/damping
    point_spacing : int; default=0
        1/2 is skip every ohter point
    use_rhoref : bool; default=False
        adds a sea level density in the user's in_units
        doesn't work at say 30,000 ft
    flutter_ncolumns : int; default=0
        number of columns for the legend; 0->dynamic
    flutter_bbox_to_anchor_x : float; default=1.0
        gap between plot and legend; 1.02 is good
    freq_ndigits : int; default=2
        precision on the frequency
    divergence_legend_loc : str; default='' -> 'best'
        The legend may overlap important info; move it
        see matplotlib documentation for options
    divergence_freq_tol : 0.05
        frequency to identify divergence (Hz)
    ndir_levels: int; default=1
        number of directory levels to write
    progress_callback : function
        run a callback for the gui

    TODO: add trade study support

    There are files of the form:
     - model_plane_mach_0.5_mgtow_kactuator_100.f06
     - model_plane_mach_0.5_mgtow_kactuator_50.f06
     - model_plane_mach_0.5_bdfw_kactuator_100.f06
     - model_plane_mach_0.5_bdfw_kactuator_50.f06
     - model_plane_mach_0.5_zfw_kactuator_100.f06
     - model_plane_mach_0.5_zfw_kactuator_50.f06

     - model_plane_mach_0.2_mgtow_kactuator_100.f06
     - model_plane_mach_0.2_mgtow_kactuator_50.f06
     - model_plane_mach_0.2_bdfw_kactuator_100.f06
     - model_plane_mach_0.2_bdfw_kactuator_50.f06
     - model_plane_mach_0.2_zfw_kactuator_100.f06
     - model_plane_mach_0.2_zfw_kactuator_50.f06

    Base              Mach   Fuel       Word   Kact
    ----------------  ----  -----  ---------   ----
    model_plane_mach  0.5   mgtow  kactuator    100
    model_plane_mach  0.5   mgtow  kactuator     50
    model_plane_mach  0.5    bdfw  kactuator    100
    model_plane_mach  0.5    bdfw  kactuator     50
    model_plane_mach  0.5     zfw  kactuator    100
    model_plane_mach  0.5     zfw  kactuator     50
    model_plane_mach  0.2   mgtow  kactuator    100
    model_plane_mach  0.2   mgtow  kactuator     50
    model_plane_mach  0.2    bdfw  kactuator    100
    model_plane_mach  0.2    bdfw  kactuator     50
    model_plane_mach  0.2     zfw  kactuator    100
    model_plane_mach  0.2     zfw  kactuator     50

    config_headers = ['Mach', 'Config', 'Kact']
    trades = [
        ['Config', 'Kact', 'Mach'],  # 1 line per (Config, Kact) with Config first
        ['Kact', 'Config', 'Mach'],  # 1 line per (Config, Kact) with Kact first
        ['Kact', 'Mach'],    # min of worst Configs
        ['Config', 'Mach'],  # min of worst Kacts
    ]
    # same as trades, but without the x-axis
    compressed_axes = [
        ['Config', 'Kact'],
        ['Kact', 'Config'],
    ]

    We want to make trade study plots. Starting with
    the yaxis, being limit, flutter, and divergence
    speeds. For the x-axis, we need a float for the
    x-axis, so options include:
     - Kact, Mach (floats)

    plot(df, case_dict, xaxis)
    """
    #-----------------------------------------
    f06_filename0 = f06_filenames[0]
    dirname = os.path.dirname(f06_filename0)
    docx_filename = os.path.join(dirname, docx_filename)
    docx_dirname = Path(docx_filename).parent
    picdir = docx_dirname / 'pics'
    if not picdir.exists():
        picdir.mkdir()

    #------------------------------
    mode_switch_method = ''  # None
    subcase = 1
    make_alt = False
    freq_round = 2
    eas_round = 1

    # The plot gets messy and dfreq_tol doesn't work if you have NaN
    # points. To hack this, we just chop the plot above some
    # x value (x_cutoff)
    x_cutoff = None  # TODO: add me; fixes NaNs?

    # point_spacing = 8
    make_pngs = True
    show_individual = False
    freq_target = 0.5
    # vl_target, vf_target
    #------------------------------
    eas_units = 'knots'  # baseline
    eas_report_units = 'KEAS'
    eas_range = eas_flutter_range
    ncol = flutter_ncolumns
    #------------------------------
    v_lines = get_vlines(vf_target, vl_target)

    if ncol in [0, 1]:
        figsize = (15, 12)
    else:
        figsize = (24, 12)

    noline, nopoints = get_noline_nopoints(
        show_lines, show_points)

    damping_crossings, damping_required_tol = _get_damping_crossings(
        damping_required, damping_required_tol,
        damping_limit)
    settings['damping_required_tol'] = str(damping_required_tol)
    settings['damping_crossings'] = str(damping_crossings)

    #------------------------------
    cases = []

    nfiles = len(f06_filenames)
    log.info(f'f06_filenames = {f06_filenames}')
    log.info(f'configs = {configs}')
    for ifile, f06_filename, config in zip(count(), f06_filenames, configs):
        log.info(f'Processing F06 {ifile}/{nfiles}: {f06_filename}')
        if progress_callback is not None:
            progress_callback(ifile, nfiles)  # 0-indexed for progress bar

        basename = os.path.splitext(os.path.basename(f06_filename))[0]
        png_filename = picdir / f'{basename}.png'
        if config.strip() == '':
            config = basename

        resp_dict, data_dict = make_flutter_response(
            str(f06_filename),
            f06_units=f06_units, out_units=out_units,
            use_rhoref=use_rhoref, make_alt=make_alt,
            log=log)
        assert len(resp_dict) == 1, resp_dict
        response = resp_dict[subcase]

        # xcutoff doesn't apply for first 6 modes
        response.nrigid_body_modes = nrigid_body_modes
        response.x_cutoff = x_cutoff

        response.noline = noline
        response.freq_ndigits = freq_ndigits

        response.set_plot_settings(
            figsize=figsize,
            # the delta spacing for the x/y axis
            # xtick_major_locator_multiple=[50., 50.],
            # ytick_major_locator_multiple=[0.02, None],
        )
        response.set_symbol_settings(
            nopoints=nopoints, show_mode_number=show_mode_number,
            point_spacing=point_spacing, markersize=5,
        )
        response.set_font_settings(plot_font_size)

        vl_vf_crossing_dict, vd_crossing_dict = response.get_flutter_crossings(
            damping_crossings=damping_crossings, modes=modes,
            eas_range=eas_range, freq_round=freq_round, eas_round=eas_round,
            divergence_freq_tol=divergence_freq_tol)

        if make_pngs:
            log.info(f"modes in plot_vg_vf = {modes}")
            fig, (damp_axes, freq_axes) = response.plot_vg_vf(
                plot_type='eas', modes=modes,
                clear=False, close=False, legend=True,
                xlim=eas_lim, ylim_damping=ylim_damping, ylim_freq=ylim_freq,
                # ivelocity: Optional[int]=None,
                v_lines=v_lines,
                damping_required=damping_required,
                damping_limit=damping_limit,
                ncol=ncol, freq_tol=freq_tol, freq_tol_remove=freq_tol_remove,
                #--------
                mode_switch_method=mode_switch_method,
                divergence_legend_loc=divergence_legend_loc,
                flutter_bbox_to_anchor=(flutter_bbox_to_anchor_x, 1.),
                # plot_freq_tol_filtered_lines=True,
                damping_crossings=damping_crossings, filter_damping=True,
                eas_range=eas_range,
                png_filename=None,
                filter_freq=True,
                show_detailed_mode_info=show_detailed_mode_info,
                show=False)

            # title = os.path.basename(png_filename)
            title = write_docx_path(f06_filename, ndir_levels=ndir_levels)
            damp_axes.set_title(title)
            # plt.tight_layout()
            # if show_individual:
            #     plt.show()
            fig.savefig(png_filename, bbox_inches='tight')
            # bbox_to_anchor=(1, 1), borderaxespad=0)
            plt.close()
        if show_individual:
            raise RuntimeError('stopping')

        # print(f'modes = {modes}')
        # print(f'xcrossing_dict = {xcrossing_dict}')
        hump_message = _get_hump_message(
            response, vl_vf_crossing_dict,
            eas_range, show_detailed_mode_info)

        log.info(f'VL_target = {vl_target}')
        v0, freq0, v3, freq3, vdiverg, freq_diverg = response.xcrossing_dict_to_VL_VF_VD(
            vl_vf_crossing_dict, vd_crossing_dict,
            log, freq_target, vl_target, vf_target,
            v_baseline=eas_max,
            # is_hump_modes=parse_hump_modes,
        )
        # if VL < VL_target:
        #     log.error(f'VL={VL} KEAS, freq={freqL} Hz; {f06_filename_base}')
        # if VF < VF_target:
        #     log.error(f'VF={VF} KEAS, freq={freqF} Hz; {f06_filename_base}')
        # if VD < VD_target:
        #     log.error(f'VD={VD} KEAS, freq={freqD} Hz; {f06_filename_base}')
        # mass = -1.0
        # cg = [0., 0., 0.]
        # inertia = [0., 0., 0., 0., 0., 0.]

        if 'opgwg' not in data_dict:
            matrices = data_dict['matrices']
            log.warning(f'data_dict_keys={list(data_dict)}; matrices_keys={list(matrices)}')
            mass = np.full(1, np.nan)
            cg = np.full(3, np.nan)
            inertia = np.full((3, 3), np.nan)
        else:
            opgwg = data_dict['opgwg']  # grid point weight
            # matrices = data_dict['matrices']
            # frequencies = matrices['freq']
            mass = opgwg['mass']
            cg = opgwg['cg']
            # print(opgwg)
            # print(f'frequencies = {frequencies.round(3)}')
            inertia = opgwg['I(S)']

        case = (v0, freq0, v3, freq3, vdiverg, freq_diverg,
                mass, cg, inertia, config,
                hump_message,
                f06_filename, png_filename)
        eas_units = response.out_units['eas']
        cases.append(case)
    _cases_to_document(
        log, docx_filename, table, cases, trades, settings,
        eas_units=eas_report_units, ndir_levels=ndir_levels)

def _cases_to_document(log: SimpleLogger,
                       docx_filename: PathLike,
                       table: pd.DataFrame,
                       cases: list,
                       trades: list[dict],
                       settings: dict[str, int | float],
                       eas_units: str='KEAS',
                       write_filename: bool=True,
                       ndir_levels: int=1):
    percent0 = settings['damping_required'] * 100
    percent3 = settings['damping_limit'] * 100
    label_vg0 = f'V,g={percent0:.0f}% ({eas_units})'
    label_vg3 = f'V,g={percent3:.0f}% ({eas_units})'
    label_vd = f'VDiverg ({eas_units})'

    label_freq_g0 = f'Freq,g={percent0:.0f}% (Hz)'
    label_freq_g3 = f'Freq,g={percent3:.0f}% (Hz)'

    configs = []
    f06_filenames = []
    # config_file_table = {
    #     'Config': configs,
    #     'File': f06_filenames,
    # }

    flutter_table = {
        # 'Configuration': [],
        'Config': configs,
        'File': f06_filenames,
        label_vg0: [],
        label_freq_g0: [],
        label_vg3: [],
        label_freq_g3: [],
        label_vd: [],
    }

    document = Document()
    _write_name_value_table(document, settings)

    for case in cases:
        # document.add_heading(f'Config={config0}', heading_level_mach)

        (v0, freq0, v3, freq3, vdiverg, freq_diverg,
         mass, cg, inertia, config,
         hump_message,
         f06_filename, png_filename) = case

        configs.append(config)
        f06_filenames.append(f06_filename)
        flutter_table[label_vg0].append(v0)
        flutter_table[label_vg3].append(v3)
        flutter_table[label_vd].append(vdiverg)
        flutter_table[label_freq_g0].append(freq0)
        flutter_table[label_freq_g3].append(freq3)

        # configi = ''
        if np.isfinite(freq0):
            v0_text = f'V 0%={v0:.0f} KEAS ({freq0:.1f} Hz)'
        else:  # TODO: why does this happen?
            v0_text = f'V 0%={v0:.0f} KEAS (default)'

        if np.isfinite(freq3):
            v3_text = f'V 3%={v3:.0f} {eas_units} ({freq3:.1f} Hz)'
        else:  # TODO: why does this happen?
            v3_text = f'V 3%={v3:.0f} {eas_units} (default)'

        text = f'{v0_text}, {v3_text}, VD={vdiverg:.0f} {eas_units}, {config}'
        if hump_message:
            text += '\n' + hump_message

        freq0_str = f'{freq0:.2f}' if np.isfinite(freq0) else 'N/A'
        freq3_str = f'{freq3:.2f}' if np.isfinite(freq3) else 'N/A'

        path_str = write_docx_path(f06_filename, ndir_levels=ndir_levels)
        if os.path.exists(png_filename):
            document.add_picture(str(png_filename), width=Inches(6.5))
        else:
            paragraph = document.add_paragraph(f'Missing {f06_filename}')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if write_filename:
            paragraph = document.add_paragraph(path_str)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # msg += f'{mach:.3f}, {weight_config}, {v0:.3f}, {freq0_str}, {v3:.3f}, {freq3_str}, {vdiverg:.3f}, {config}, {f06_filename_base}\n'

    if percent0 <= -100.0:
        del flutter_table[label_vg0]
        del flutter_table[label_freq_g0]
    if percent3 <= -100.0:
        del flutter_table[label_vg3]
        del flutter_table[label_freq_g3]
    _write_2d_table(document, flutter_table, log, 'Flutter Results')

    if len(trades) and 0:
        codestr = 'from envelope import plot_mach_eas'
        # module = importlib.import_module(module_name)

        # Then, use getattr to get the specific class/function
        # DequeClass = getattr(module, class_name)

        # trades_expected = [
        #     (
        #         ['3', '5', '2'],
        #         {('bdfw', 50): [3, 7],
        #          ('mgtow', 50): [1, 5],
        #          ('bdfw', 100): [2, 6],
        #          ('mgtow', 100): [0, 4]},
        #     )
        # ]
        fig = plt.figure()
        ax = fig.gca()
        # flutter_table = {
        #     # 'Configuration': [],
        #     'Config': configs,
        #     'File': f06_filenames,
        #     label_vg0: [],
        #     label_freq_g0: [],
        #     label_vg3: [],
        #     label_freq_g3: [],
        #     label_vd: [],
        # }
        for names, trade_dict in trades:
            v0s = flutter_table.get(label_vg0, None)
            vfs = flutter_table.get(label_vg3, None)
            vds = flutter_table.get(label_vd, None)
            *dep_vars, ind_var = names
            ind_values = table[ind_var]
            def plot(*args):
                pass
            plot(ax, trade_dict,
                 ind_var, ind_values,
                 v0s, vfs, vds)
    log.info(f'saving docx {docx_filename}')
    document.save(docx_filename)
    return
    # ncol = flutter_ncolumns
    #
    # damping_crossings = [
    #     (0.00, 0.01),
    #     (0.03, 0.03),
    # ]
    #
    # damping_crossings2 = {}
    # if isinstance(damping_crossings, list):
    #     for key, value in damping_crossings:
    #         damping_crossings2[key] = value
    #     damping_crossings = damping_crossings2
    #     damping_crossings2 = {}
    #
    # for key, value in damping_crossings.items():
    #     if np.allclose(key, 0.0):
    #         damping_crossings2[key] = key + 0.001
    #     damping_crossings2[key] = value
    #
    # if filename_base_to_vl_vf_vbase_dict is None:
    #     filename_base_to_vl_vf_vbase_dict = {}
    # assert isinstance(filename_base_to_vl_vf_vbase_dict, dict), filename_base_to_vl_vf_vbase_dict
    # print(f'filename_base_to_vl_vf_vbase_dict = {filename_base_to_vl_vf_vbase_dict}')
    #
    # if case_format == 'func':
    #     assert case_info_func is not None, case_info_func
    #
    # if x_cutoff is None:
    #     x_cutoff = eas_lim[1]
    #
    # # type checking
    # str(eas_lim[0] + 1)
    # str(eas_lim[1] + 1)
    #
    # if sub_dirnames is None:
    #     log.warning(f'sub_dirnames = {sub_dirnames}; assuming single folder')
    #     sub_dirnames = [None]
    # if skip_extra_cases:
    #     log.error(f'skip_extra_cases={skip_extra_cases}; should be temporary')
    #
    # # ---------------------------------------------------------------------------------
    # cases = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
    # cases2 = defaultdict(list)
    # for sub_dirname in sub_dirnames:
    #     sub_dirname2: str = '' if sub_dirname is None else sub_dirname
    #     # if sub_dirname is None:
    #     #     dirnamei = dirname
    #     # else:
    #     dirnamei = dirname / sub_dirname2
    #
    #     dirnamei_pics = dirname / f'pics_{sub_dirname}'
    #     print(dirnamei_pics)
    #     # if not dirnamei_pics.exists():
    #     os.makedirs(dirnamei_pics, exist_ok=True)
    #
    #     print(f'directory = {str(dirnamei)}')
    #     print(f'glob_str = {str(glob_str)}')
    #     # f06_root_filenames = [fname for fname in os.listdir(dirnamei) if fname.endswith('.f06')]
    #     f06_root_filenames = dirnamei.glob('*.f06')
    #     f06_root_filenames = natsorted(f06_root_filenames)
    #     f06_root_filenames = filter_f06(f06_root_filenames, glob_str)
    #     # f06_filenames = [dirnamei / fname for fname in f06_root_filenames]
    #     f06_filenames = f06_root_filenames
    #
    #     for f06_filename in f06_filenames:
    #         print('------------------------------------------------')
    #         assert os.path.exists(f06_filename), print_bad_path(f06_filename)
    #
    #         f06_filename_base = f06_filename.name
    #
    #         base1 = os.path.splitext(f06_filename_base)[0]
    #         png_filename = dirnamei_pics / (base1 + '.png')
    #         print(f'f06_filename = {str(f06_filename)}')
    #         # print(f'png_filename = {str(png_filename)}')
    #         is_failed, out = get_case_info(
    #             f06_filename_base,
    #             store_configs,
    #             weight_configs,
    #             machs_to_parse, machs_to_skip,
    #             case_format, case_info_func, skip_extra_cases)
    #         if is_failed:
    #             continue
    #         mach_num_str = out['mach_num_str']
    #         mach_number = out['mach']
    #         weight_config = out['weight_config']
    #         configi = out['config']
    #         # mach_num_str, mach_number, weight_config, configi = out
    #
    #         log.level = 'debug'
    #         # print(f'f06_filename_base = {f06_filenme_base}')
    #         try:
    #             resp_dict, data_dict = make_flutter_response(
    #                 str(f06_filename),
    #                 f06_units='english_in', out_units='english_kt',
    #                 use_rhoref=use_rhoref,
    #                 log=log)
    #             assert len(resp_dict) == 1, resp_dict
    #         except:
    #             raise
    #             if stop_on_failure:
    #                 raise
    #             VL = -1.0
    #             VF = -1.0
    #             VD = -1.0
    #             freqL = -1.0
    #             freqF = -1.0
    #             # freqD = -1.0
    #             mass = np.zeros(1)[0]
    #             cg = np.zeros(3)
    #             inertia = np.zeros((3, 3))
    #             png_filename = ''
    #             case_key = (mach_number, weight_config, store_config)
    #             case_value = (VL, freqL, VF, freqF, VD,
    #                           mass, cg, inertia, configi,
    #                           '',
    #                           f06_filename_base, f06_filename, png_filename)
    #             cases2[case_key].append(case_value)
    #             cases[mach_number][weight_config][store_config].append(case_value)
    #             continue
    #
    #         resp = resp_dict[1]
    #         modes = resp.modes[6:]
    #         assert modes[0] == 7, modes
    #
    #         # xcutoff doesn't apply for first 6 modes
    #         resp.nrigid_body_modes = 6
    #         resp.x_cutoff = x_cutoff
    #         resp.set_plot_settings(
    #             figsize=figsize,
    #             xtick_major_locator_multiple=[50., 50.],
    #             ytick_major_locator_multiple=[0.02, None],
    #         )
    #         resp.set_symbol_settings(
    #             nopoints=False, show_mode_number=False, point_spacing=8,
    #             markersize=5,
    #         )
    #
    #         # Crossing = tuple[float, float, float]
    #         # if parse_hump_modes:
    #         #     vl_vf_crossing_dict, hump_vd_crossing_dict = resp.get_hump_flutter_crossings(
    #         #         damping_crossings=damping_crossings, modes=modes,
    #         #         eas_range=eas_range)
    #         #     print(hump_vd_crossing_dict)
    #         #     vl_vf_crossing_dict = hump_vd_crossing_dict
    #         # else:
    #         vl_vf_crossing_dict, vd_crossing_dict = resp.get_flutter_crossings(
    #             damping_crossings=damping_crossings2, modes=modes,
    #             eas_range=eas_range)
    #         # print(vl_vf_crossing_dict)
    #
    #         if make_pngs:
    #             fig, (damp_axes, freq_axes) = resp.plot_vg_vf(
    #                 plot_type='eas', modes=modes,
    #                 clear=False, close=False, legend=True,
    #                 xlim=eas_lim, ylim_damping=ylim_damping, ylim_freq=ylim_freq,
    #                 # ivelocity: Optional[int]=None,
    #                 v_lines=v_lines,
    #                 damping_required=0.0,
    #                 damping_limit=0.03,
    #                 ncol=ncol, freq_tol=freq_tol, freq_tol_remove=freq_tol_remove,
    #                 # plot_freq_tol_filtered_lines=True,
    #                 damping_crossings=damping_crossings2, filter_damping=True,
    #                 eas_range=eas_range,
    #                 png_filename=None,
    #                 filter_freq=True,
    #                 show_detailed_mode_info=show_detailed_mode_info,
    #                 show=False)
    #             basename = os.path.basename(png_filename)
    #             damp_axes.set_title(basename)
    #             # plt.tight_layout()
    #             # if show_individual:
    #             #     plt.show()
    #             fig.savefig(png_filename, bbox_inches='tight')
    #             # bbox_to_anchor=(1, 1), borderaxespad=0)
    #             shutil.copyfile(png_filename, png_filename_mach)
    #             plt.close()
    #         if show_individual:
    #             raise RuntimeError('stopping')
    #
    #         if show_detailed_mode_info and 0:
    #             keys = list(filename_base_to_vl_vf_vbase_dict)
    #             # print('filename_base_to_vl_vf_vbase_dict.keys = ', keys)
    #             # print(f'f06_filename_base = {f06_filename_base}')
    #             if len(filename_base_to_vl_vf_vbase_dict):
    #                 v0i, vfi, vbase = filename_base_to_vl_vf_vbase_dict[f06_filename_base]
    #                 print(f'v0i={v0i}, vfi={vfi}, vbase={vbase}')
    #                 vl_array, vf_array = resp.hump_modes_from_VL_VF_dict(
    #                     vl_vf_crossing_dict, v0i, vfi, vbase, log)
    #                 if len(vl_array):
    #                     print('vl:')
    #                     print(vl_array)
    #                 if len(vf_array):
    #                     print('vf:')
    #                     print(vf_array)
    #
    #         # print(f'modes = {modes}')
    #         # print(f'xcrossing_dict = {xcrossing_dict}')
    #         hump_message = _get_hump_message(
    #             resp, vl_vf_crossing_dict,
    #             eas_range, show_detailed_mode_info)
    #
    #         log.info(f'VL_target = {VL_target}')
    #         VL, freqL, VF, freqF, VD, freqD = resp.xcrossing_dict_to_VL_VF_VD(
    #             vl_vf_crossing_dict, vd_crossing_dict,
    #             log, freq_target, VL_target, VF_target,
    #             v_baseline=V_baseline,
    #             # is_hump_modes=parse_hump_modes,
    #         )
    #         if VL < VL_target:
    #             log.error(f'VL={VL} KEAS, freq={freqL} Hz; {f06_filename_base}')
    #         if VF < VF_target:
    #             log.error(f'VF={VF} KEAS, freq={freqF} Hz; {f06_filename_base}')
    #         if VD < VD_target:
    #             log.error(f'VD={VD} KEAS, freq={freqD} Hz; {f06_filename_base}')
    #
    #         case_key = (mach_number, weight_config)
    #         case_value = (VL, freqL, VF, freqF, VD,
    #                       mass, cg, inertia, configi,
    #                       hump_message,
    #                       f06_filename_base, f06_filename, png_filename)
    #         cases2[case_key].append(case_value)
    #         cases[mach_number][weight_config].append(case_value)
    #         log.debug(f'mach={mach_number} VL={VL:g} VF={VF:g} VD={VD:g} {f06_filename_base}')
    #         if hump_message:
    #             log.warning(hump_message)
    #             # asdf
    #         # print(xcrossing_dict)
    #         # return cases
    #
    # names = ['Mach', 'Weight', 'Config']
    # case_data = CaseData(cases, cases2, names)
    # return case_data

def _get_hump_message(resp: FlutterResponse,
                      vl_vf_crossing_dict,
                      eas_range: tuple[float, float],
                      show_detailed_mode_info: bool) -> str:
    hump_message = ''
    if not show_detailed_mode_info:
        return hump_message

    hump_message_list = resp.get_hump_mode_messages(
        vl_vf_crossing_dict,
        modes=None,
        eas_range=eas_range,
        filter_damping=False,
        plot_type='eas',
    )
    if hump_message_list:
        hump_message_list2 = [line.strip().replace('\n', '; ') for line in hump_message_list]
        hump_message = '\n'.join(hump_message_list2)
        # log.warning(f'hump_message = {hump_message!r}')
    return hump_message


def _write_name_value_table(document, records: dict[str, float]) -> None:
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Value'
    for name, value in records.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(name)
        row_cells[1].text = str(value)


def _write_2d_table(document: Document,
                    records: dict[str, list[list[float]]],
                    log: SimpleLogger, table_name: str) -> None:
    """
    Expects a 2D table (e.g., list of lists)

    Parameters
    ----------
    document : Document
        the word doc
    records: dict[str, list[float]
        the data to add

    Returns
    -------
    mass_table = {
        'Configuration': ['Clean'],
        'Mass (in)': [1.0],
        'XCG (in)': [1.0],
        'YCG (in)': [1.0],
        'ZCG (in)': [1.0],
        'Ixx (slinch-in^2)': [1.0],
        'Iyy (slinch-in^2)': [1.0],
        'Izz (slinch-in^2)': [1.0],
    }
    """
    keys = list(records.keys())
    nkeys = len(keys)
    assert nkeys > 0, keys
    key0 = keys[0]
    column0 = records[key0]
    ncols = nkeys
    nrows = len(column0) + 1

    table = document.add_table(rows=nrows, cols=ncols)

    # write the column headers
    hdr_cells = table.rows[0].cells
    for ikey, key in enumerate(keys):
        hdr_cells[ikey].text = key

    # write the rows
    icol = 0
    for key, values in records.items():
        for irow in range(1, nrows):
            # print(key, irow, nrows, values)
            try:
                value = values[irow-1]
            except IndexError:
                log.error(f'problem writing {table_name} for irow={irow}')
                continue
            row_cells = table.rows[irow].cells
            row_cells[icol].text = str(value)
        icol += 1
    return
