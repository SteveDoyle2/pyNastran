"""
TODO: add default directory so when an f06 is bad, you don't
      have to keep changing the directory
"""
import os
import sys
import copy
import warnings
import traceback
from itertools import count
from pathlib import Path
from functools import wraps
from typing import Optional, Any
import natsort

import numpy as np
from matplotlib import pyplot as plt
import matplotlib.gridspec as gridspec

from pyNastran.utils import print_bad_path, PathLike
from pyNastran.utils.dev import get_files_of_type

try:
    import json5 as json
except ModuleNotFoundError:
    warnings.warn('couldnt find json5, using json')
    import json

from pyNastran.f06.dev.flutter.utils import get_raw_json
JSON_FILENAME, USE_VTK, USE_TABS = get_raw_json(allow_vtk=False)

from qtpy import QtCore
Qt = QtCore.Qt
from qtpy.compat import getopenfilename  # getsavefilename
# from qtpy.QtGui import QIcon, QPixmap
from qtpy.QtWidgets import (
    QLabel, QWidget,
    QApplication, QVBoxLayout, QComboBox,  # QMenu, QLineEdit,
    QHBoxLayout, QPushButton, QGridLayout,
    QAction,
    QCheckBox, QLineEdit,
    QListWidgetItem, QAbstractItemView,
    QListWidget, QSpinBox, QTabWidget,  # QToolButton,
    QTableWidget, QTableWidgetItem, QMenu, QInputDialog,
    QFileDialog, QProgressBar,
)
# from qtpy.QtWidgets import (
#     QMessageBox,
#     QMainWindow, QDockWidget, QFrame, QToolBar,
#     QToolButton, QMenuBar,
# )
# from qtpy.QtGui import QIcon

from cpylog import SimpleLogger
import pyNastran
from pyNastran.gui.utils.qt.pydialog import QFloatEdit, make_font
from pyNastran.gui.qt_files.named_dock_widget import NamedDockWidget
from pyNastran.gui.qt_files.loggable_gui import LoggableGui

from pyNastran.f06.dev.flutter.gui_flutter import (
    export_flutter_results, get_list_float_or_none, get_float_or_none,
    get_selected_items_flat,
)

from pyNastran.f06.dev.flutter.actions_builder import Actions, Action, build_menus
from pyNastran.f06.dev.flutter.preferences_object import FlutterPreferencesObject
from pyNastran.f06.dev.flutter.preferences import (
    FLUTTER_BBOX_TO_ANCHOR_DEFAULT, LEGEND_LOC_DEFAULT,
    FONT_SIZE_DEFAULT, FLUTTER_NCOLUMNS_DEFAULT, FREQ_NDIGITS_DEFAULT, FREQ_DIVERGENCE_TOL)

from pyNastran.f06.flutter_response import Limit  # FlutterResponse
from pyNastran.f06.parse_flutter import get_flutter_units

from pyNastran.f06.dev.flutter.utils import (
    validate_json,
    get_point_removal_str,
    point_removal_str_to_point_removal,
    _float_passed_to_default, _to_str,
    get_plot_flags, get_plot_file,
    update_ylog_style, get_png_filename,
    load_f06_op2, get_vlines, get_damping_crossings,
    X_PLOT_TYPES, PLOT_TYPES, UNITS_IN, UNITS_OUT,
    MODE_SWITCH_METHODS)

PKG_PATH = Path(pyNastran.__path__[0])

AERO_PATH = PKG_PATH / '..' / 'models' / 'aero'

from pyNastran.f06.dev.flutter.vtk_data import VtkData
import pandas as pd
import tables

from pyNastran.f06.parse_flutter import make_flutter_response, FlutterResponse
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    warnings.warn('>>> pip install python-docx')
    raise

QLINEEDIT_WHITE = 'QLineEdit {background-color: white;}'
QLINEEDIT_RED = 'QLineEdit {background-color: red;}'
ICON_PATH = Path('')


def split_by_pattern(strings: list[str],
                     delimiter: str='_',
                     group_common: bool=True):
    """
    Main function to split strings based on common pattern.
    """
    if not strings:
        return []

    # Split all strings
    split_lists = [s.split(delimiter) for s in strings]
    if not group_common or len(split_lists) < 2:
        return split_lists
    split_list0 = split_lists[0]
    # print(f'split_list0 = {split_list0}')

    # Find common prefix
    common_prefix_len = 0
    min_length = min(len(parts) for parts in split_lists)

    for i in range(min_length):
        if all(parts[i] == split_lists[0][i] for parts in split_lists):
            common_prefix_len += 1
            continue
        break
    # print(f'common_prefix_len = {common_prefix_len}')
    prefix = split_list0[:common_prefix_len]
    # print(f'prefix = {prefix}')

    # Find common suffix
    common_suffix_len = 0
    for i in range(1, min_length - common_prefix_len + 1):
        if all(parts[-i] == split_lists[0][-i] for parts in split_lists):
            common_suffix_len += 1
            continue
        break
    print(f'common_suffix_len = {common_suffix_len}')

    # Reconstruct
    result = []
    for parts in split_lists:
        new_parts = []
        if common_prefix_len > 0:
            new_parts.append(delimiter.join(parts[:common_prefix_len]))

        middle_start = common_prefix_len
        middle_end = len(parts) - common_suffix_len if common_suffix_len > 0 else len(parts)
        new_parts.extend(parts[middle_start:middle_end])

        if common_suffix_len > 0:
            new_parts.append(delimiter.join(parts[-common_suffix_len:]))
        result.append(new_parts)
    return result

class QTableWidgetCopy(QTableWidget):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # self.setContextMenuPolicy(Qt.CustomContextMenu)
        # self.customContextMenuRequested.connect(self.show_context_menu)

        rename_column_support = True
        add_remove_row_support = True

        if rename_column_support:
            # Enable custom context menu for horizontal header
            self.horizontalHeader().setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
            self.horizontalHeader().customContextMenuRequested.connect(self.show_header_context_menu)

        if add_remove_row_support:
            # Add/remove rows
            self.verticalHeader().setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
            self.verticalHeader().customContextMenuRequested.connect(self.show_row_header_context_menu)

    def load_table_data(self,
                        headers: list[str],
                        data: np.ndarray | list[list[int | float | str]]) -> None:
        """
        headers = ['A', 'Col2', 'Col3', 'Col4']
        data = np.arange(20).reshape(4,5)
        self.load_table_data(headers, data)
        """
        # Set row and column counts
        # print(f'data = {data}')
        self.clear()
        self.setRowCount(len(data))
        self.setColumnCount(len(data[0]) if len(data) else 0)

        # Set headers (optional)
        # headers = ["Column 1", "Column 2", "Column 3"]  # Replace with your headers
        self.setHorizontalHeaderLabels(headers)

        # Populate the table
        for row_idx, row_data in enumerate(data):
            for col_idx, col_data in enumerate(row_data):
                item = QTableWidgetItem(str(col_data))
                self.setItem(row_idx, col_idx, item)

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_C and (event.modifiers() & Qt.KeyboardModifier.ControlModifier):
            self.on_copy_table(event)

    def on_copy_table(self, event):
        copied_cells = self.selectedIndexes()
        if not copied_cells:
            return
        # Sort to maintain order, handle rows/cols with \t and \n
        # copied_cells = sorted(copied_cells)

        copy_text = ''
        max_column = copied_cells[-1].column()
        for c in copied_cells:
            copy_text += self.item(c.row(), c.column()).text()
            if c.column() == max_column:
                copy_text += '\n'
            else:
                copy_text += '\t'
        QApplication.clipboard().setText(copy_text)

    def show_context_menu(self, pos):
        if self.rowCount() == 0 or self.columnCount() == 0:
            return
        print(f'pos = {pos}')
        # Get the item at the clicked position
        item = self.itemAt(pos)
        print(f'item = {item}')
        if not item:
            return
        # Get global position to display menu correctly
        global_pos = self.mapToGlobal(pos)
        print(f'global_pos = {global_pos}')

        # Create the menu
        context_menu = QMenu(self)
        copy_action = context_menu.addAction('Copy')
        delete_action = context_menu.addAction('Delete Row')

        # Execute the menu and wait for an action selection
        action = context_menu.exec_(global_pos)

        if action == copy_action:
            # Handle copy logic
            print(f'Copying cell content: {item.text()}')
        elif action == delete_action:
            # Handle delete logic
            row = item.row()
            self.removeRow(row)
            # print(f'Deleting row {row}')

    def show_header_context_menu(self, position):
        """Show context menu when right-clicking on column header"""
        # Get the column index from the position
        column = self.horizontalHeader().logicalIndexAt(position)
        if column < 0:
            return

        # Create context menu
        menu = QMenu(self)
        rename_action = menu.addAction("Rename Column")

        # copyAction = menu.addAction('Copy')
        delete_action = menu.addAction('Delete Column')
        insert_left_action = menu.addAction('Insert Column Left')
        insert_right_action = menu.addAction('Insert Column Right')

        # Show menu and get selected action
        action = menu.exec(self.horizontalHeader().mapToGlobal(position))

        if action == rename_action:
            self.rename_column(column)
        elif action == insert_left_action:
            col = self.horizontalHeader().logicalIndexAt(position)
            self.insert_col(col)
        elif action == insert_right_action:
            col = self.horizontalHeader().logicalIndexAt(position)
            self.insert_col(col+1)
        elif action == delete_action:
            col = self.horizontalHeader().logicalIndexAt(position)
            self.delete_col(col)

    def rename_column(self, column: int):
        """Open dialog to rename column header"""
        # Get current column name
        current_name = self.horizontalHeaderItem(column).text() if self.horizontalHeaderItem(
            column) else f"Column {column}"

        # Show input dialog
        new_name, ok = QInputDialog.getText(
            self,
            "Rename Column",
            "Enter new column name:",
            text=current_name
        )

        if ok and new_name:
            self.setHorizontalHeaderItem(column, QTableWidgetItem(new_name))
    #-----------------
    # def show_col_header_context_menu(self, position):
    #     """Show context menu when right-clicking on row header (row numbers)"""
    #     # Get the col index from the position
    #     col = self.horizontalHeader().logicalIndexAt(position)
    #
    #     if col < 0:
    #         return
    #
    #     # Create context menu
    #     menu = QMenu(self)
    #     insert_left_action = menu.addAction("Insert Column Left")
    #     insert_right_action = menu.addAction("Insert Column Right")
    #     menu.addSeparator()
    #     delete_action = menu.addAction("Delete Col")
    #
    #     # Show menu and get selected action
    #     action = menu.exec(self.horizontalHeader().mapToGlobal(position))
    #
    #     if action == insert_left_action:
    #         self.insert_col(col)
    #     elif action == insert_right_action:
    #         self.insert_col(col + 1)
    #     elif action == delete_action:
    #         self.delete_col(col)

    def show_row_header_context_menu(self, position):
        """Show context menu when right-clicking on row header (row numbers)"""
        # Get the row index from the position
        row = self.verticalHeader().logicalIndexAt(position)

        if row < 0:
            return

        # Create context menu
        menu = QMenu(self)
        insert_above_action = menu.addAction("Insert Row Above")
        insert_below_action = menu.addAction("Insert Row Below")
        menu.addSeparator()
        delete_action = menu.addAction("Delete Row")

        # Show menu and get selected action
        action = menu.exec(self.verticalHeader().mapToGlobal(position))

        if action == insert_above_action:
            self.insert_row(row)
        elif action == insert_below_action:
            self.insert_row(row + 1)
        elif action == delete_action:
            self.delete_row(row)

    def insert_row(self, row: int):
        """Insert a new empty row at the specified position"""
        self.insertRow(row)

        # Initialize empty cells in the new row
        for col in range(self.columnCount()):
            self.setItem(row, col, QTableWidgetItem(""))
        # print(f'Inserted row at position {row}')

    def insert_col(self, col: int):
        """Insert a new empty col at the specified position"""
        self.insertColumn(col)

        # Initialize empty cells in the new row
        for col in range(self.columnCount()):
            self.setItem(col, col, QTableWidgetItem(""))
        # print(f'Inserted col at position {col}')

    def delete_row(self, row: int):
        """Delete the specified row"""
        if self.rowCount() > 0:
            self.removeRow(row)
            # print(f'Deleted row {row}')

    def delete_col(self, col: int):
        """Delete the specified col"""
        if self.columnCount() > 0:
            self.removeColumn(col)
            # print(f'Deleted col {col}')

    def delete_selected_rows(self):
        """Delete all selected rows (triggered by Delete key)"""
        selected_rows = set()
        for item in self.selectedItems():
            selected_rows.add(item.row())

        # Delete from highest to lowest to avoid index shifting issues
        for row in sorted(selected_rows, reverse=True):
            self.removeRow(row)

        if selected_rows:
            print(f'Deleted {len(selected_rows)} row(s)')

    def delete_selected_cols(self):
        """Delete all selected rows (triggered by Delete key)"""
        selected_cols = set()
        for item in self.selectedItems():
            selected_cols.add(item.row())

        # Delete from highest to lowest to avoid index shifting issues
        for col in sorted(selected_cols, reverse=True):
            self.removeColumn(col)

        if selected_cols:
            print(f'Deleted {len(selected_cols)} col(s)')

    def get_data(self,
                 skip_empty_rows=True,
                 strip_whitespace=True,
                 convert_numeric=False):
            """
            Extract data from QTableWidget to pandas DataFrame with advanced options

            Parameters
            ----------
            skip_empty_rows : bool
                If True, skip rows that are entirely empty (default: True)
            strip_whitespace : bool
                If True, strip leading/trailing whitespace from all cells (default: True)
            convert_numeric : bool
                If True, attempt to convert numeric strings to numbers (default: False)

            Returns
            -------
            pandas.DataFrame
                DataFrame containing the table data
            """
            rows = self.rowCount()
            cols = self.columnCount()

            # Get headers
            headers = []
            for col in range(cols):
                header = self.horizontalHeaderItem(col)
                if header is not None:
                    headers.append(header.text())
                else:
                    headers.append(f"Column_{col}")
            # print(f'headers = {headers}')

            # Get data
            data = []
            for row in range(rows):
                row_data = []
                for col in range(cols):
                    item = self.item(row, col)
                    if item is not None:
                        cell_value = item.text()
                        if strip_whitespace:
                            cell_value = cell_value.strip()
                        row_data.append(cell_value)
                    else:
                        row_data.append("")
                # print(f'row_data = {row_data}')

                # Check if row is entirely empty
                if skip_empty_rows:
                    if any(cell for cell in row_data):  # If any cell has content
                        data.append(row_data)
                else:
                    data.append(row_data)

            # Convert numeric columns if requested
            df = pd.DataFrame(data, columns=headers)
            if convert_numeric:
                df = df.apply(pd.to_numeric, errors='ignore')
            return df


class FlutterGui(LoggableGui):
    def __init__(self, f06_filename: str=''):
        super().__init__(html_logging=True)
        self.use_vtk = USE_VTK
        self.use_tabs = USE_TABS
        self.vtk_data = VtkData()

        self._export_settings_obj = FlutterPreferencesObject(self, USE_VTK, hide_vtk=True)
        self.iwindows = []

        self.divergence_legend_loc = LEGEND_LOC_DEFAULT
        self.flutter_bbox_to_anchor_x = FLUTTER_BBOX_TO_ANCHOR_DEFAULT
        self.flutter_ncolumns = FLUTTER_NCOLUMNS_DEFAULT
        self.freq_ndigits = FREQ_NDIGITS_DEFAULT
        self.freq_divergence_tol = FREQ_DIVERGENCE_TOL
        self.auto_update = True

        # Mach	Fuel	Filename
        # 0.2	0	mach0.2_fuel_0.f06
        # 0.3	40	mach0.3_fuel_40.f06
        # 0.4	100	mach0.4_fuel_100.f06
        self.excel_filename = r'C:\work\test.xlsx'
        self.font_size = FONT_SIZE_DEFAULT
        self.plot_font_size = FONT_SIZE_DEFAULT
        self.show_lines = True
        self.show_points = True
        self.show_mode_number = False
        self.show_detailed_mode_info = False
        self.point_spacing = 0
        self.use_rhoref = False
        self._units_in = ''
        self._units_out = ''
        self.units_in = ''
        self.units_out = ''
        self.use_dock_widgets = self.html_logging
        self.qactions = {}
        self.nrecent_files_max = 20
        self.recent_files = []
        self.save_filename = JSON_FILENAME
        self.is_valid = False

        self.data = {}
        self.f06_filename = ''
        self.bdf_filename = ''
        self.op2_filename = ''
        self._bdf_filename_default = ''
        self._op2_filename_default = ''
        self.subcase = 0
        self.x_plot_type = 'eas'
        self.plot_type = 'x-damp-freq'
        self.mode_switch_method = MODE_SWITCH_METHODS[0]
        self.point_removal = ''

        self.eas_lim = []
        self.tas_lim = []
        self.mach_lim = []
        self.alt_lim = []
        self.q_lim = []
        self.rho_lim = []
        self.eas_flutter_range = [None, None]
        # self.eas_diverg_range = [None, None]
        self.freq_lim = [None, None]
        # self.damping_lim = [None, None]
        self.ydamp_lim = [None, None]
        self.kfreq_lim = [None, None]
        self.eigr_lim = [None, None]
        self.eigi_lim = [None, None]
        self.responses = {}
        self.modes = []
        self.selected_modes = []
        self.freq_tol = -1.0
        self.freq_tol_remove = -1.0
        self.mag_tol = -1.0
        self.damping = -1.0
        self.damping_required = -1.0
        self.damping_required_tol = -1.0
        self.vf = -1.0
        self.vl = -1.0
        self.export_to_png = True
        self.export_to_csv = False
        self.export_to_f06 = False
        self.export_to_zaero = False

        self.setup_lists()
        self.setup_widgets()
        self.setup_layout()
        self.on_load_settings()
        ifile = -1
        if f06_filename:
            self.f06_filename_edit[ifile].setText(f06_filename)
            self._set_f06_default_names(f06_filename)

        self.setup_toolbar()
        self._update_recent_files_actions()
        self.setup_connections()
        self._set_window_title()
        self.on_font_size()
        self.on_plot_type()
        self.setAcceptDrops(True)
        # self.on_open_new_window()
        self.show()

    def dragEnterEvent(self, event) -> None:
        if event.mimeData().hasUrls():
            event.accept()
        else:
            event.ignore()

    def dropEvent(self, event) -> None:
        ifile = self.ifile
        filenames = [url.toLocalFile() for url in event.mimeData().urls()]
        for filename in filenames:
            flower = filename.lower()
            if flower.endswith(('.f06', '.out')):
                self.f06_filename_edit[ifile].setText(filename)
            else:
                self.log.error(f'unknown extension (f06) format for {filename}')

    def setup_toolbar(self) -> None:
        # frame = QFrame(self)
        actions_dict = {
            # 'file_load': Action(name='file_load', text='Load...', func=self.on_file_load, icon='folder.png'),
            # 'file_save': Action(name='file_save', text='Save...', func=self.on_file_save, icon='save.png'),
            # 'file_save_as': Action(name='file_save_as', text='Save As...', func=self.on_file_save_as),
            'file_exit':       Action(name='exit', text='Exit...', icon='exit2.jpg', func=self.on_file_exit),
            'export_settings': Action(name='Export Settings', text='Export Settings...', icon='preferences.jpg',
                                      shortcut='Ctrl+P', func=self.on_export_settings),
        }
        actions_input = Actions(ICON_PATH, actions_dict)  # load_icon=False
        recent_files = actions_input.build_recent_file_qactions(
            self, self.recent_files, self.set_f06)
        self.qactions = actions_input.build_qactions(self)

        file_actions = [
            # 'file_load',
            # 'file_save', 'file_save_as',
            ] + recent_files + [
            'file_exit']
        view_actions = ['export_settings']

        self.menubar = self.menuBar()
        self.file_menu = self.menubar.addMenu('File')
        self.view_menu = self.menubar.addMenu('View')
        # self.help_menu = self.menubar.addMenu('Help')

        # help_actions = []
        menus_dict = {
            'File': (self.file_menu, file_actions),
            'View': (self.view_menu, view_actions),
            # 'Help': (self.help_menu, help_actions),
        }
        build_menus(menus_dict, self.qactions)
        # self.file_menu.addAction(actions['file_load'])
        # self.file_menu.addAction(actions['exit'])

        # self.toolbar = self.addToolBar('Show toolbar')
        # self.toolbar.setObjectName('main_toolbar')
        self.statusbar = self.statusBar()

    def dont_crash(func):
        @wraps(func)
        def wrapper(self, *args, **kwargs):
            # do something before `sum`
            try:
                result = func(self, *args, **kwargs)
            except Exception as e:
                self.log_error(str(traceback.format_exc()))
                result = None
            # do something after `sum`
            return result
        return wrapper

    def set_f06(self, ifile: int) -> None:
        f06_filename = self.recent_files[ifile]
        self.f06_filename_edit[self.ifile].setText(f06_filename)
        self.on_load_f06(None)

    def setup_modes(self) -> None:
        self.modes_widget = QListWidget(self)
        self.modes_widget.setMaximumWidth(200)  # was 100 when no freq
        self.modes_widget.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self._set_modes_table(self.modes_widget, [0], [0.])

    def dontcrash(func):
        @wraps(func)
        def wrapper(self, *args, **kwargs):
            # do something before `sum`
            result = func(self) # , *args, **kwargs
            # do something after `sum`
            return result
        return wrapper

    # @dontcrash
    def on_export_settings(self) -> None:
        self._export_settings_obj.show()

    def on_file_exit(self) -> None:
        if hasattr(self, 'on_file_save') and hasattr(self, 'save_filename'):
            self.on_file_save()

    def on_file_save(self) -> None:
        if self.save_filename == '' or not os.path.exists(self.save_filename):
            self.on_file_save_as()
        else:
            # self.log.warning('on_file_save; _save')
            self._save(self.save_filename)

    def on_file_save_as(self) -> None:
        # if 0:
        #     #print('on_file_save')
        #     qt_wildcard = '*.json'
        #     basedir = str(DIRNAME)
        #     json_filename, wildcard = getsavefilename(
        #         self, caption=title, basedir=basedir,
        #         filters=qt_wildcard,
        #         #options=QFileDialog.setLabelText('data.json'),
        #     )
        #     self.log.info(f'json_filename={json_filename!r} wildcard={wildcard!r}')
        json_filename = self.save_filename
        # self.log.warning('on_file_save_as; _save')
        self._save(json_filename)

    def _save(self, json_filename: str) -> None:
        is_valid = self.validate()
        # self.log.info(f'self.data = {self.data}; is_valid={is_valid}')
        if json_filename == '' or len(self.data) == 0:
            return
        # print(f'json_filename={json_filename!r} wildcard={wildcard!r}')
        # print(f'self.data = {self.data}')
        out_data = copy.deepcopy(self.data)
        out_data['use_vtk'] = self.use_vtk
        # out_data['use_tabs'] = False
        if USE_VTK:
            out_data['vtk'] = self._vtk_window_obj.data
        out_data['preferences'] = {
            'flutter_ncolumns': self.flutter_ncolumns,
            'freq_ndigits': self.freq_ndigits,
            'freq_divergence_tol': self.freq_divergence_tol,
            'auto_update': self.auto_update,
            'flutter_bbox_to_anchor_x': self.flutter_bbox_to_anchor_x,
            'divergence_legend_loc': self.divergence_legend_loc,
            'export_to_png': self.export_to_png,
            'export_to_csv': self.export_to_csv,
            'export_to_f06': self.export_to_f06,
            'export_to_zaero': self.export_to_zaero,
            'font_size': self.font_size,
            'plot_font_size': self.plot_font_size,
        }
        with open(json_filename, 'w') as json_file:
            json.dump(out_data, json_file, indent=4)
        # print(f'fname="{fname}"')
        self.log.info(f'finished saving {json_filename!r}\n')
        # self.log.info(f'out_data = {out_data}; is_valid={is_valid}')
        self.save_filename = json_filename
        self._set_window_title()

    def _apply_settings(self, data: dict[str, Any]) -> None:
        if USE_VTK:
            self.vtk_data.apply_settings(data)
            self._vtk_window_obj.apply_settings(data)
        log = self.log
        font_size0 = self.font_size
        # radios = [
        #     ('show_points', self.show_points_radio),
        # ]
        # for (key, checkbox) in radios:
        #     if key not in data:
        #         continue
        #     val = data[key]
        #     assert isinstance(val, bool), (key, val)
        #     checkbox.setChecked(val)

        # preferences = data.get('preferences', {})
        # for key, value in preferences.items():
        #     if not hasattr(self, key):
        #         self.log.error(f'load failure: skipping {key!r}={value!r} because key doesnt exist')
        #         continue
        #     self.log.error(f'load success: loaded {key!r}={value!r}')
        #     setattr(self, key, value)

        spinners = [
             # ('plot_font_size', self.plot_font_size_edit),
        ]
        for (key, spinner) in spinners:
            if key not in data:
                continue
            val = data[key]
            assert isinstance(val, int), (key, val)
            spinner.setValue(val)

        type_names = [
            (int,  ('preferences/font_size',
                    'preferences/plot_font_size',
                    'preferences/flutter_ncolumns')),
            (float, ('preferences/flutter_bbox_to_anchor_x',)),
            (str, ('preferences/divergence_legend_loc',)),
            (bool, ('preferences/export_to_png',
                    'preferences/export_to_f06',
                    'preferences/export_to_csv',
                    'preferences/export_to_zaero')),
        ]
        for value_type, keys in type_names:
            for key in keys:
                # print(f'loading key={key!r}')
                if '/' in key:
                    skey = key.split('/')
                    assert len(skey) == 2, f'key={key!r} skey={skey}'
                    key0, key1 = skey
                    if key0 not in data:
                        log.warning(f'skipping {key!r} because {key0} does not exist')
                        print(data)
                        continue
                    data0 = data[key0]
                    value = data0[key1]
                    assert hasattr(self, key1), (key, value)
                    if not isinstance(value, value_type):
                        log.warning(f'{key!r}={value!r} and is not {value_type}...skipping')
                        continue
                    # log.info(f'setting {key!r} (key1={key1!r}) -> {value!r}')
                    setattr(self, key1, value)
                else:
                    if key not in data:
                        log.warning(f'skipping {key!r}')
                        assert len(key) > 1, keys
                        continue
                    value = data[key]
                    assert hasattr(self, key), (key, value)
                    if not isinstance(value, value_type):
                        log.warning(f'{key!r}={value!r} and is not {value_type}...skipping')
                        continue
                    # log.info(f'setting {key!r} -> {value!r}')
                    setattr(self, key, value)

        ifile = self.ifile
        checkboxs = [
            ('use_rhoref', self.use_rhoref_checkbox),
            ('show_points', self.show_points_checkbox[ifile]),
            ('show_mode_number', self.show_mode_number_checkbox[ifile]),
            ('show_detailed_mode_info', self.show_detailed_mode_info_checkbox[ifile]),
            ('show_lines', self.show_lines_checkbox[ifile]),
        ]
        # attrs aren't stored
        for (key, checkbox) in checkboxs:
            if key not in data:
                continue
            val = data[key]
            assert isinstance(val, bool), (key, val)
            try:
                checkbox.setChecked(val)
            except AttributeError:  # pragma: no cover
                print(key)
                raise

        min_max_line_edits = [
            ('eas_lim', self.eas_lim_edit_min, self.eas_lim_edit_max),
            ('tas_lim', self.tas_lim_edit_min, self.tas_lim_edit_max),
            ('mach_lim', self.mach_lim_edit_min, self.mach_lim_edit_max),
            ('alt_lim', self.alt_lim_edit_min, self.alt_lim_edit_max),
            ('q_lim', self.q_lim_edit_min, self.q_lim_edit_max),
            ('rho_lim', self.rho_lim_edit_min, self.rho_lim_edit_max),
            ('ikfreq_lim', self.ikfreq_lim_edit_min, self.ikfreq_lim_edit_max),

            ('damp_lim', self.damp_lim_edit_min, self.damp_lim_edit_max),
            ('eas_flutter_range', self.eas_flutter_range_edit_min, self.eas_flutter_range_edit_max),
            # ('eas_diverg_range', self.eas_diverg_range_edit_min, self.eas_diverg_range_edit_max),
            ('freq_lim', self.freq_lim_edit_min, self.freq_lim_edit_max),
            ('kfreq_lim', self.kfreq_lim_edit_min, self.kfreq_lim_edit_max),
        ]
        for key, line_edit_min, line_edit_max in min_max_line_edits:
            if key not in data:
                # print(f'apply_settings: skipping key={key!r}')
                continue
            values = data[key]
            value0 = _to_str(values[0])
            value1 = _to_str(values[1])
            line_edit_min.setText(value0)
            line_edit_max.setText(value1)

        point_removal = data.get('point_removal', [])
        point_removal_str = get_point_removal_str(point_removal)
        self.point_removal_edit.setText(point_removal_str)

        line_edits = [
            ('recent_files', 0, self.f06_filename_edit[ifile]),
            ('freq_tol', -1, self.freq_tol_edit),
            ('freq_tol_remove', -1, self.freq_tol_remove_edit),
            ('mag_tol', -1, self.mag_tol_edit),
            ('vl', -1, self.VL_edit),
            ('vf', -1, self.VF_edit),
            ('damping', -1, self.damping_edit),
            ('damping_required', -1, self.damping_required_edit),
            ('damping_required_tol', -1, self.damping_required_tol_edit),
            ('output_directory', -1, self.output_directory_edit),
        ]
        for key, index, line_edit in line_edits:
            if key not in data:
                # print(f'apply_settings: skipping key={key!r}')
                continue
            values = data[key]
            if index != -1:
                value = values[index]
            else:
                value = values

            str_value = _to_str(value)

            # print('type(value) =', type(value))
            # print(f'{key+":":<10} values={values}[{index!r}]={value!r} -> {str_value!r}')
            try:
                line_edit.setText(str_value)
            except AttributeError:  # pragma: no cover
                print(key)
                raise

        pulldown_edits = [
            ('x_plot_type', self.x_plot_type_pulldown, X_PLOT_TYPES),
            ('plot_type', self.plot_type_pulldown, PLOT_TYPES),
            ('units_in', self.units_in_pulldown, UNITS_IN),
            ('units_out', self.units_out_pulldown, UNITS_OUT),
            ('mode_switch_method', self.mode_switch_method_pulldown, MODE_SWITCH_METHODS),
        ]
        for key, pulldown_edit, values in pulldown_edits:
            if key not in data:
                # print(f'apply_settings: skipping key={key!r}')
                continue
            value = data[key]
            index = values.index(value)
            pulldown_edit.setCurrentIndex(index)

        self.recent_files = []
        for fname in data['recent_files']:
            abs_path = os.path.abspath(fname)
            if abs_path not in self.recent_files:
                self.recent_files.append(abs_path)
        self.f06_filename = self.recent_files[0]
        if self.font_size != font_size0:
            self.on_set_font_size(self.font_size)

    def on_enable_bdf(self) -> None:
        state = self.bdf_filename_checkbox.isChecked()
        self.bdf_filename_edit.setEnabled(state)
        self.bdf_filename_browse.setEnabled(state)
        if state:
            self.bdf_filename_edit.setText(self.bdf_filename)
        else:
            self.bdf_filename_edit.setText(self._bdf_filename_default)

    def on_enable_op2(self) -> None:
        state = self.op2_filename_checkbox.isChecked()
        self.op2_filename_edit.setEnabled(state)
        self.op2_filename_browse.setEnabled(state)
        if state:
            self.op2_filename_edit.setText(self.op2_filename)
        else:
            self.op2_filename_edit.setText(self._op2_filename_default)

    def on_browse_f06(self) -> None:
        """pops a dialog to select the f06 file"""
        title = 'Load a Flutter (Nastran F06, Zona Out) File'
        qt_wildcard = 'F06 File (*.f06);; Zona File (*.out)'
        basedir = os.path.dirname(self.f06_filename)
        fname, wildcard_level = getopenfilename(
            self, caption=title, basedir=basedir, filters=qt_wildcard,)
        if fname == '':
            return
        self.f06_filename_edit[self.ifile].setText(fname)
        self.run_button.setEnabled(False)

    # @dontcrash
    def on_load_settings(self) -> None:
        json_filename = self.save_filename
        if not os.path.exists(json_filename):
            self.log.warning(f'unable to find {json_filename}')
            return
        # self.save_filename = json_filename
        try:
            with open(json_filename, 'r') as json_file:
                data = json.load(json_file)
            is_valid = validate_json(data, self.log)
            self._apply_settings(data)
        except Exception as e:
            self.log.error(f'failed to load {json_filename}\n{str(e)}')
            print(traceback.format_exc())
            # print(traceback.format_exception_only(e))
            # raise
            return
        self.log.info(f'finished loading {json_filename!r}')
        # return wildcard_level, fname
        self._set_window_title()

    def _set_window_title(self) -> None:
        if self.save_filename == '':
            self.setWindowTitle('Flutter Plot')
        else:
            self.setWindowTitle(f'Flutter Plot: {self.save_filename}')

    def setup_lists(self) -> None:
        self.f06_filename_label = []
        self.f06_filename_edit = []
        self.f06_filename_browse = []

        self.bdf_filename_checkbox = []
        self.bdf_filename_edit = []
        self.bdf_filename_browse = []

        self.op2_filename_checkbox = []
        self.op2_filename_edit = []
        self.op2_filename_browse = []

        # self.use_rhoref_checkbox = []
        #
        # self.log_xscale_checkbox = []
        # self.log_yscale1_checkbox = []
        # self.log_yscale2_checkbox = []

        self.show_points_checkbox = []
        self.show_mode_number_checkbox = []
        self.show_detailed_mode_info_checkbox = []
        self.point_spacing_label = []
        self.point_spacing_spinner = []
        self.show_lines_checkbox = []

        self.include_rigid_body_modes_checkbox = []
        self.number_rigid_body_modes_label = []
        self.number_rigid_body_modes_spinner = []

        # self.index_lim_label = []
        # self.index_lim_edit_min = []
        # self.index_lim_edit_max = []
        #
        # self.eas_lim_label = []
        # self.eas_lim_edit_min = []
        # self.eas_lim_edit_max = []
        #
        # self.tas_lim_label = []
        # self.tas_lim_edit_min = []
        # self.tas_lim_edit_max = []
        #
        # self.mach_lim_label = []
        # self.mach_lim_edit_min = []
        # self.mach_lim_edit_max = []
        #
        # self.alt_lim_label = []
        # self.alt_lim_edit_min = []
        # self.alt_lim_edit_max = []
        #
        # self.q_lim_label = []
        # self.q_lim_edit_min = []
        # self.q_lim_edit_max = []
        #
        # self.rho_lim_label = []
        # self.rho_lim_edit_min = []
        # self.rho_lim_edit_max = []
        #
        # self.damp_lim_label = []
        # self.damp_lim_edit_min = []
        # self.damp_lim_edit_max = []
        #
        # self.freq_lim_label = []
        # self.freq_lim_edit_min = []
        # self.freq_lim_edit_max = []
        #
        # self.kfreq_lim_label = []
        # self.kfreq_lim_edit_min = []
        # self.kfreq_lim_edit_max = []
        #
        # self.ikfreq_lim_label = []
        # self.ikfreq_lim_edit_min = []
        # self.ikfreq_lim_edit_max = []
        #
        # self.eigr_lim_label = []
        # self.eigr_lim_edit_min = []
        # self.eigr_lim_edit_max = []
        #
        # self.eigi_lim_label = []
        # self.eigi_lim_edit_min = []
        # self.eigi_lim_edit_max = []
        #
        # #--------------------------------------------
        # self.freq_tol_label = []
        # self.freq_tol_edit = []
        #
        # self.mag_tol_label = []
        # self.mag_tol_edit = []
        #
        # self.freq_tol_remove_label = []
        # self.freq_tol_remove_edit = []
        #
        # self.subcase_label = []
        # self.subcase_edit = []
        # self.x_plot_type_label = []
        # self.x_plot_type_pulldown = []
        #
        # self.plot_type_label = []
        # self.plot_type_pulldown = []
        #
        # self.units_in_label = []
        # self.units_in_pulldown = []
        #
        # self.units_out_label = []
        # self.units_out_pulldown = []
        #
        # self.output_directory_label = []
        # self.output_directory_edit = []
        # self.output_directory_browse = []
        #
        # self.VL_label = []
        # self.VL_edit = []
        #
        # self.VF_label = []
        # self.VF_edit = []
        #
        # self.damping_label = []
        # self.damping_edit = []
        #
        # self.eas_flutter_range_label = []
        # self.eas_flutter_range_edit_min = []
        # self.eas_flutter_range_edit_max = []
        #
        # self.eas_diverg_range_label = []
        # self.eas_diverg_range_edit_min = []
        # self.eas_diverg_range_edit_max = []
        #
        # self.point_removal_label = []
        # self.point_removal_edit = []
        #
        # self.mode_label = []
        # self.mode_edit = []
        #
        # self.velocity_label = []
        # self.velocity_edit = []
        #
        # self.f06_load_button = []
        # self.run_button = []
        #
        # self.pop_vtk_gui_button = []

        # self.solution_type_label = []
        # self.solution_type_pulldown = []
        # self.mode2_label = []
        # self.mode2_pulldown = []
        #
        # self.mode_switch_method_label = []
        # self.mode_switch_method_pulldown = []

    def setup_widgets(self) -> None:
        self.f06_filename_label.append(QLabel('F06 Filename:', self))
        self.f06_filename_edit.append(QLineEdit(self))
        self.f06_filename_browse.append(QPushButton('Browse...', self))

        self.bdf_filename_checkbox = QCheckBox('BDF Filename:', self)
        self.bdf_filename_edit = QLineEdit(self)
        self.bdf_filename_browse = QPushButton('Browse...', self)
        self.bdf_filename_checkbox.setChecked(False)
        self.bdf_filename_edit.setEnabled(False)
        self.bdf_filename_browse.setEnabled(False)
        self.bdf_filename_edit.setToolTip('Loads the Nastran Geometry')

        self.op2_filename_checkbox = QCheckBox('OP2 Filename:', self)
        self.op2_filename_edit = QLineEdit(self)
        self.op2_filename_browse = QPushButton('Browse...', self)
        self.op2_filename_checkbox.setChecked(False)
        self.op2_filename_edit.setEnabled(False)
        self.op2_filename_browse.setEnabled(False)
        self.op2_filename_edit.setToolTip('Loads the Nastran Results (and geometry if BDF Filename is empty)')

        self.use_rhoref_checkbox = QCheckBox('Sea Level Rho Ref', self)
        self.use_rhoref_checkbox.setChecked(False)

        self.log_xscale_checkbox = QCheckBox('Log Scale x', self)
        self.log_yscale1_checkbox = QCheckBox('Log Scale y1', self)
        self.log_yscale2_checkbox = QCheckBox('Log Scale y2', self)
        self.log_xscale_checkbox.setChecked(False)
        self.log_yscale1_checkbox.setChecked(False)
        self.log_yscale2_checkbox.setChecked(False)

        self.show_points_checkbox.append(QCheckBox('Show Points', self))
        self.show_mode_number_checkbox.append(QCheckBox('Show Mode Number', self))
        self.show_detailed_mode_info_checkbox.append(QCheckBox('Show Detailed Mode Info', self))
        self.point_spacing_label.append(QLabel('Point Spacing', self))
        self.point_spacing_spinner.append(QSpinBox(self))
        self.include_rigid_body_modes_checkbox.append(QCheckBox('Include Rigid Body Modes', self))
        self.number_rigid_body_modes_label.append(QLabel('nRigid Body Modes', self))
        self.number_rigid_body_modes_spinner.append(QSpinBox(self))

        for obj in self.include_rigid_body_modes_checkbox:
            obj.setVisible(False)
        for obj in self.number_rigid_body_modes_label:
            obj.setVisible(False)
        for obj in self.number_rigid_body_modes_spinner:
            obj.setVisible(False)

        self.show_lines_checkbox.append(QCheckBox('Show Lines', self))
        self.show_points_checkbox[-1].setChecked(True)
        self.show_lines_checkbox[-1].setChecked(True)
        self.show_points_checkbox[-1].setToolTip('The points are symbols')
        self.show_mode_number_checkbox[-1].setToolTip('The points are the mode number')
        self.show_detailed_mode_info_checkbox[-1].setToolTip('Lists the 0% eas/freq range')
        self.point_spacing_spinner[-1].setToolTip('Skip Every Nth Point; 0=Plot All')
        self.point_spacing_spinner[-1].setValue(0)
        self.point_spacing_spinner[-1].setMinimum(0)
        self.point_spacing_spinner[-1].setMaximum(30)

        self.index_lim_label = QLabel('Index Limits:', self)
        self.index_lim_edit_min = QFloatEdit('0', self)
        self.index_lim_edit_max = QFloatEdit(self)

        self.eas_lim_label = QLabel('EAS Limits:', self)
        self.eas_lim_edit_min = QFloatEdit('0', self)
        self.eas_lim_edit_max = QFloatEdit(self)

        self.tas_lim_label = QLabel('TAS Limits:', self)
        self.tas_lim_edit_min = QFloatEdit('0', self)
        self.tas_lim_edit_max = QFloatEdit(self)

        self.mach_lim_label = QLabel('Mach Limits:', self)
        self.mach_lim_edit_min = QFloatEdit(self)
        self.mach_lim_edit_max = QFloatEdit(self)

        self.alt_lim_label = QLabel('Alt Limits:', self)
        self.alt_lim_edit_min = QFloatEdit(self)
        self.alt_lim_edit_max = QFloatEdit(self)

        self.q_lim_label = QLabel('Q Limits:', self)
        self.q_lim_edit_min = QFloatEdit(self)
        self.q_lim_edit_max = QFloatEdit(self)

        self.rho_lim_label = QLabel('Rho Limits:', self)
        self.rho_lim_edit_min = QFloatEdit('0', self)
        self.rho_lim_edit_max = QFloatEdit(self)

        self.damp_lim_label = QLabel('Damping Limits (g):', self)
        self.damp_lim_edit_min = QFloatEdit('-0.3', self)
        self.damp_lim_edit_max = QFloatEdit('0.3', self)

        self.freq_lim_label = QLabel('Freq Limits (Hz):', self)
        self.freq_lim_edit_min = QFloatEdit('0', self)
        self.freq_lim_edit_max = QFloatEdit(self)

        self.kfreq_lim_label = QLabel('KFreq Limits:', self)
        self.kfreq_lim_edit_min = QFloatEdit(self)
        self.kfreq_lim_edit_max = QFloatEdit(self)

        self.ikfreq_lim_label = QLabel('1/KFreq Limits:', self)
        self.ikfreq_lim_edit_min = QFloatEdit(self)
        self.ikfreq_lim_edit_max = QFloatEdit(self)

        self.eigr_lim_label = QLabel('Real Eigenvalue:', self)
        self.eigr_lim_edit_min = QFloatEdit(self)
        self.eigr_lim_edit_max = QFloatEdit(self)

        self.eigi_lim_label = QLabel('Imag Eigenvalue:', self)
        self.eigi_lim_edit_min = QFloatEdit(self)
        self.eigi_lim_edit_max = QFloatEdit(self)

        # --------------------------------------------
        self.freq_tol_label = QLabel('dFreq Tol (Hz) Dash:', self)
        self.freq_tol_edit = QFloatEdit('-1.0', self)
        self.freq_tol_edit.setToolTip("Applies a dotted line for modes that don't change by more than some amount")

        self.freq_tol_remove_label = QLabel('dFreq Tol (Hz) Remove:', self)
        self.freq_tol_remove_edit = QFloatEdit('-1.0', self)
        self.freq_tol_remove_edit.setToolTip('Removes a mode if it meets dFreq Tol (Hz) Dash and Remove')

        self.mag_tol_label = QLabel('Magnitude Tol:', self)
        self.mag_tol_edit = QFloatEdit('-1.0', self)
        self.mag_tol_edit.setToolTip('Filters modal participation factors based on magnitude')

        self.subcase_label = QLabel('Subcase:', self)
        self.subcase_edit = QComboBox(self)

        units_msg = (
            "english_in: inch/s, slich/in^3\n"
            "english_ft: ft/s,   slug/ft^3\n"
            "english_kt: knots,  slug/ft^3\n"
            "si:         m/s,    kg/m^3\n"
            "si-mm:      mm/s,   Mg/mm^3\n"
        )
        self.x_plot_type_label = QLabel('X-Axis Plot Type:', self)
        self.x_plot_type_pulldown = QComboBox(self)
        self.x_plot_type_pulldown.addItems(X_PLOT_TYPES)
        self.x_plot_type_pulldown.setToolTip('sets the x-axis')

        self.plot_type_label = QLabel('Plot Type:', self)
        self.plot_type_pulldown = QComboBox(self)
        self.plot_type_pulldown.addItems(PLOT_TYPES)
        # self.plot_type_pulldown.setToolTip(units_msg)

        self.units_in_label = QLabel('Units In:', self)
        self.units_in_pulldown = QComboBox(self)
        self.units_in_pulldown.addItems(UNITS_IN)
        self.units_in_pulldown.setToolTip(units_msg)
        iunits_in = UNITS_IN.index('english_in')
        self.units_in_pulldown.setCurrentIndex(iunits_in)
        self.units_in_pulldown.setToolTip('Sets the units for the F06/OP2; set when loaded')

        self.units_out_label = QLabel('Units Out:', self)
        self.units_out_pulldown = QComboBox(self)
        self.units_out_pulldown.addItems(UNITS_OUT)
        self.units_out_pulldown.setToolTip(units_msg)
        iunits_out = UNITS_IN.index('english_kt')
        self.units_out_pulldown.setCurrentIndex(iunits_out)
        self.units_out_pulldown.setToolTip('Sets the units for the plot; may be updated')

        self.output_directory_label = QLabel('Output Directory:', self)
        self.output_directory_edit = QLineEdit('', self)
        self.output_directory_browse = QPushButton('Browse...', self)
        self.output_directory_edit.setDisabled(True)
        self.output_directory_browse.setDisabled(True)

        self.VL_label = QLabel('VL, Limit:', self)
        self.VL_edit = QFloatEdit('', self)
        self.VL_edit.setToolTip('Makes a vertical line for VL')

        self.VF_label = QLabel('VF, Flutter:', self)
        self.VF_edit = QFloatEdit('', self)
        self.VF_edit.setToolTip('Makes a vertical line for VF')

        self.damping_required_label = QLabel('Damping Required, g:', self)
        self.damping_required_edit = QFloatEdit('', self)
        self.damping_required_edit.setToolTip('Enables the flutter crossing (e.g., 0.0 for 0%)')

        self.damping_required_tol_label = QLabel('Damping Required Tol, g:', self)
        self.damping_required_tol_edit = QFloatEdit('', self)
        self.damping_required_tol_edit.setToolTip('Tolerance for Damping Required. The crossing will be reported at the required value')

        self.damping_label = QLabel('Damping, g:', self)
        self.damping_edit = QFloatEdit('', self)
        self.damping_edit.setToolTip('Enables the flutter crossing (e.g., 0.03 for 3%)')

        self.eas_flutter_range_label = QLabel('EAS Flutter/Diverg Range:', self)
        self.eas_flutter_range_edit_min = QFloatEdit('', self)
        self.eas_flutter_range_edit_max = QFloatEdit('', self)
        self.eas_flutter_range_edit_min.setToolTip('Defines the flutter/divergence crossing range')
        self.eas_flutter_range_edit_max.setToolTip('Defines the flutter/divergence crossing range')

        # self.eas_diverg_range_label = QLabel('EAS Diverg Range:', self)
        # self.eas_diverg_range_edit_min = QFloatEdit('', self)
        # self.eas_diverg_range_edit_max = QFloatEdit('', self)
        # self.eas_diverg_range_edit_min.setToolTip('Defines the divergence crossing range')
        # self.eas_diverg_range_edit_max.setToolTip('Defines the divergence crossing range')

        self.point_removal_label = QLabel('Point Removal:', self)
        self.point_removal_edit = QLineEdit('', self)
        self.point_removal_edit.setToolTip('Remove bad points from a mode; "400:410,450:500"')

        self.mode_label = QLabel('Mode:', self)
        self.mode_edit = QSpinBox(self)
        self.mode_edit.setMinimum(1)
        # self.mode_edit.SetValue(3)
        self.mode_edit.setToolTip('Sets the mode')

        self.velocity_label = QLabel('Velocity Point:', self)
        self.velocity_edit = QComboBox(self)
        self.velocity_edit.setToolTip('Sets the velocity (input units)')

        self.f06_load_button = QPushButton('Load F06', self)
        self.run_button = QPushButton('Run', self)

        self.pop_vtk_gui_button = QPushButton('Open GUI', self)
        self.solution_type_label = QLabel('Solution Type:', self)
        self.solution_type_pulldown = QComboBox(self)
        self.mode2_label = QLabel('Mode:', self)
        self.mode2_pulldown = QComboBox(self)

        self.mode_switch_method_label = QLabel('Mode Switch Method:', self)
        self.mode_switch_method_pulldown = QComboBox(self)
        self.mode_switch_method_pulldown.addItems(MODE_SWITCH_METHODS)

        self.setup_modes()
        self.on_plot_type()
        self.on_enable_bdf()
        self.on_enable_op2()
        self.on_hide_vtk()

    def on_hide_vtk(self) -> None:
        if USE_VTK:
            return
        objs = [
            self.bdf_filename_checkbox, self.bdf_filename_edit, self.bdf_filename_browse,
            self.op2_filename_checkbox, self.op2_filename_edit, self.op2_filename_browse,
            self.pop_vtk_gui_button, self.solution_type_label, self.solution_type_pulldown,
            self.mode2_label, self.mode2_pulldown,
        ]
        for obj in objs:
            obj.setVisible(False)

    def on_plot_type(self) -> None:
        x_plot_type = self.x_plot_type_pulldown.currentText()
        plot_type = self.plot_type_pulldown.currentText()
        self.on_units_out()

        flags = get_plot_flags(plot_type, x_plot_type)
        show_index_lim = flags['show_index_lim']
        show_eas_lim = flags['show_eas_lim']
        show_tas_lim = flags['show_tas_lim']
        show_mach_lim = flags['show_mach_lim']
        show_alt_lim = flags['show_alt_lim']
        show_q_lim = flags['show_q_lim']
        show_rho_lim = flags['show_rho_lim']

        show_xlim = flags['show_xlim']
        show_freq = flags['show_freq']
        show_damp = flags['show_damp']
        show_kfreq = flags['show_kfreq']
        show_ikfreq = flags['show_ikfreq']
        show_root_locus = flags['show_root_locus']
        show_zimmerman = flags['show_zimmerman']
        show_modal_participation = flags['show_modal_participation']

        # print(f'x_plot_type={x_plot_type} show_damp={show_damp}; show_xlim={show_xlim}')
        # assert show_xlim is False, show_xlim

        show_eigenvalue = show_root_locus or show_modal_participation
        show_xaxis = not show_eigenvalue
        show_freq_tol = show_xaxis or show_root_locus
        show_crossing = show_xaxis
        show_mode_switch = show_xaxis or show_root_locus

        self.mode_switch_method_label.setVisible(show_mode_switch)
        self.mode_switch_method_pulldown.setVisible(show_mode_switch)

        self.x_plot_type_label.setVisible(show_xaxis)
        self.x_plot_type_pulldown.setVisible(show_xaxis)
        self.freq_tol_label.setVisible(show_freq_tol)
        self.freq_tol_edit.setVisible(show_freq_tol)
        self.freq_tol_remove_edit.setVisible(show_freq_tol)

        self.index_lim_label.setVisible(show_index_lim)
        self.index_lim_edit_min.setVisible(show_index_lim)
        self.index_lim_edit_max.setVisible(show_index_lim)

        self.eas_lim_label.setVisible(show_eas_lim)
        self.eas_lim_edit_min.setVisible(show_eas_lim)
        self.eas_lim_edit_max.setVisible(show_eas_lim)

        self.tas_lim_label.setVisible(show_tas_lim)
        self.tas_lim_edit_min.setVisible(show_tas_lim)
        self.tas_lim_edit_max.setVisible(show_tas_lim)

        self.mach_lim_label.setVisible(show_mach_lim)
        self.mach_lim_edit_min.setVisible(show_mach_lim)
        self.mach_lim_edit_max.setVisible(show_mach_lim)

        self.alt_lim_label.setVisible(show_alt_lim)
        self.alt_lim_edit_min.setVisible(show_alt_lim)
        self.alt_lim_edit_max.setVisible(show_alt_lim)

        self.q_lim_label.setVisible(show_q_lim)
        self.q_lim_edit_min.setVisible(show_q_lim)
        self.q_lim_edit_max.setVisible(show_q_lim)

        self.rho_lim_label.setVisible(show_rho_lim)
        self.rho_lim_edit_min.setVisible(show_rho_lim)
        self.rho_lim_edit_max.setVisible(show_rho_lim)

        self.damp_lim_label.setVisible(show_damp)
        self.damp_lim_edit_min.setVisible(show_damp)
        self.damp_lim_edit_max.setVisible(show_damp)

        self.damping_required_label.setVisible(show_crossing)
        self.damping_required_edit.setVisible(show_crossing)
        self.damping_required_tol_label.setVisible(show_crossing)
        self.damping_required_tol_edit.setVisible(show_crossing)

        self.damping_label.setVisible(show_crossing)
        self.damping_edit.setVisible(show_crossing)

        self.freq_lim_label.setVisible(show_freq)
        self.freq_lim_edit_min.setVisible(show_freq)
        self.freq_lim_edit_max.setVisible(show_freq)

        self.kfreq_lim_label.setVisible(show_kfreq)
        self.kfreq_lim_edit_min.setVisible(show_kfreq)
        self.kfreq_lim_edit_max.setVisible(show_kfreq)

        self.ikfreq_lim_label.setVisible(show_ikfreq)
        self.ikfreq_lim_edit_min.setVisible(show_ikfreq)
        self.ikfreq_lim_edit_max.setVisible(show_ikfreq)

        self.eigr_lim_label.setVisible(show_root_locus)
        self.eigr_lim_edit_min.setVisible(show_root_locus)
        self.eigr_lim_edit_max.setVisible(show_root_locus)

        self.eigi_lim_label.setVisible(show_root_locus)
        self.eigi_lim_edit_min.setVisible(show_root_locus)
        self.eigi_lim_edit_max.setVisible(show_root_locus)

        self.VL_label.setVisible(show_eas_lim)
        self.VL_edit.setVisible(show_eas_lim)
        self.VF_label.setVisible(show_eas_lim)
        self.VF_edit.setVisible(show_eas_lim)
        # ifile = self.ifile
        ifile = -1

        show_items = [
            (show_modal_participation, (
                self.mode_label, self.mode_edit,
                self.velocity_label, self.velocity_edit,
                self.mag_tol_label, self.mag_tol_edit)),
            (not show_modal_participation, (
                self.point_spacing_label[ifile], self.point_spacing_spinner[ifile],
                self.show_mode_number_checkbox[ifile],
                self.show_detailed_mode_info_checkbox[ifile],
                self.show_lines_checkbox[ifile],
                self.log_xscale_checkbox,
                self.log_yscale1_checkbox, self.log_yscale2_checkbox,
            ),),
        ]
        for show_hide, items in show_items:
            for item in items:
                item.setVisible(show_hide)

    def setup_layout(self) -> None:
        self.tabs = QTabWidget()
        tab_file = QWidget()
        tab_org = QWidget()
        iwindow = self.tabs.addTab(tab_file, 'File')
        iwindow_organize = self.tabs.addTab(tab_org, 'Organize')
        self.iwindows = [iwindow, iwindow_organize]

        vbox_file = self._setup_file_layout()
        vbox_org = self._setup_organize_layout()
        tab_file.setLayout(vbox_file)
        tab_org.setLayout(vbox_org)
        self.setCentralWidget(self.tabs)
        tab_file.activateWindow()

    def _setup_organize_layout(self) -> None:
        self.excel_filename_label = QLabel('Excel Filename:', self)
        self.excel_filename_edit = QLineEdit(self)
        self.excel_filename_edit.setText(self.excel_filename)
        self.excel_filename_edit.setToolTip('Must click load button before selecting tab')
        self.excel_filename_browse = QPushButton('Browse...', self)

        self.base_f06_directory = ''
        self.base_f06_directory = r'C:\work\code\pyNastran\models\aero\2_mode_flutter\dev'
        # self.base_f06_directory = r'C:\NASA\m4\formats\git\pyNastran\models\aero\2_mode_flutter'
        self.base_f06_directory_label = QLabel('Base F06 Directory:', self)
        self.base_f06_directory_edit = QLineEdit(self)
        # self.base_f06_directory_edit.setText(self.base_f06_directory)
        self.base_f06_directory_browse = QPushButton('Browse...', self)
        self.base_f06_directory_load = QPushButton('Load...', self)

        self.word_filename = 'flutter_summary.docx'
        self.word_filename_label = QLabel('Word Filename:', self)
        self.word_filename_edit = QLineEdit(self)
        self.word_filename_edit.setText(self.word_filename)
        self.word_filename_browse = QPushButton('Browse...', self)

        self.load_excel_button = QPushButton('Load...')
        self.tab_select_pulldown = QComboBox(self)

        self.tab_label = QLabel('Tab Name:', self)
        self.tab_edit = QLineEdit(self)
        self.tab_edit.setVisible(False)

        # self.tab_select_browse.setEnabled(False)
        # self.tab_select_browse.setToolTip('Must click load button before selecting tab')

        self.run_organize_button = QPushButton('Run', self)

        # self.starting_row = QLabel('Starting Row:', self)
        # self.starting_row = QLineEdit(self)
        #
        # self.starting_col = QLabel('Starting Col:', self)
        # self.starting_col = QLineEdit(self)

        self.table_widget = QTableWidgetCopy(self)
        self.progress_bar = QProgressBar(self)
        self.progress_bar.setVisible(False)

        #---------------------------------------------------
        grid = QGridLayout()

        file_row = 1
        grid.addWidget(self.base_f06_directory_label, file_row, 0)
        grid.addWidget(self.base_f06_directory_edit, file_row, 1)
        grid.addWidget(self.base_f06_directory_browse, file_row, 2)
        grid.addWidget(self.base_f06_directory_load, file_row, 3)

        file_row += 1
        grid.addWidget(self.excel_filename_label, file_row, 0)
        grid.addWidget(self.excel_filename_edit, file_row, 1)
        grid.addWidget(self.excel_filename_browse, file_row, 2)
        grid.addWidget(self.load_excel_button, file_row, 3)

        file_row += 1
        grid.addWidget(self.tab_label, file_row, 0)
        # grid.addWidget(self.tab_edit, file_row, 1)
        grid.addWidget(self.tab_select_pulldown, file_row, 1)

        grid2 = QGridLayout()
        file_row = 1
        grid2.addWidget(self.word_filename_label, file_row, 0)
        grid2.addWidget(self.word_filename_edit, file_row, 1)
        grid2.addWidget(self.word_filename_browse, file_row, 2)

        vbox = QVBoxLayout()
        vbox.addLayout(grid)
        # vbox.addWidget(self.load_excel_button)
        vbox.addWidget(self.table_widget)
        vbox.addLayout(grid2)
        vbox.addWidget(self.run_organize_button)
        vbox.addWidget(self.progress_bar)

        self.load_excel_button.clicked.connect(self.on_load_excel)
        self.tab_select_pulldown.currentIndexChanged.connect(self.on_select_excel_tab)

        # self.excel_filename_browse.clicked.connect(self.on_select_excel_tab)
        self.base_f06_directory_load.clicked.connect(self.on_base_f06_directory_load)
        self.base_f06_directory_browse.clicked.connect(self.on_base_f06_directory_browse)
        self.excel_filename_browse.clicked.connect(self.on_load_excel_file)
        self.word_filename_browse.clicked.connect(self.on_load_word_file)

        # self.base_f06_directory_browse.setEnabled(False)
        self.tab_select_pulldown.setEnabled(False)
        # self.run_organize_button.setEnabled(False)
        # self.excel_filename_browse.setEnabled(False)
        self.run_organize_button.clicked.connect(self.on_run_organize)

        self.base_f06_directory_edit.setText(self.base_f06_directory)
        return vbox

    @dontcrash
    def on_run_organize(self) -> None:
        is_valid = self.validate()
        log = self.log
        if not is_valid:
            return

        #----------------------------------------------
        configs = []
        out_table = self.table_widget.get_data()
        # print(f'out_table:\n{out_table}')

        columns = out_table.columns
        config_headers_lower = [column.lower().strip() for column in columns]
        config_headers = [column.strip() for column in columns if 'config' in config_headers_lower]
        if 'config' in config_headers_lower:
            iconfig_key = config_headers_lower.index('config')
            configs = out_table[iconfig_key].to_list()
            config_headers.remove('config')
        else:
            log.error('missing Config from case table')

        if 'Filename' not in columns:
            log.error('missing Filename from case table')
            return
        #----------------------------------------------
        # configs = [config if config.strip() else for config in ]

        word_filename = self.word_filename_edit.text().strip()
        if len(word_filename) == 0:
            word_filename = 'flutter_summary.docx'
        if not word_filename.lower().endswith('.docx'):
            word_filename += 'docx'

        # docx_filename = dirname / f'{prefix}flutter.docx'
        # from pyNastran.utils import object_attributes
        # print(object_attributes(self))

        print(f'data = {self.data}')
        # vl = self.vl
        # vl = self.data['vl']
        # data = {
        #     # 'bdf_filename': self.bdf_filename,
        #     # 'op2_filename': self.op2_filename,
        #     'log_scale_x': self.log_xscale_checkbox.isChecked(),
        #     'log_scale_y1': self.log_yscale1_checkbox.isChecked(),
        #     'log_scale_y2': self.log_yscale2_checkbox.isChecked(),
        #     'use_rhoref': self.use_rhoref,
        #     'show_points': self.show_points,
        #     'show_mode_number': self.show_mode_number,
        #     'show_detailed_mode_info': self.show_detailed_mode_info,
        #     'point_spacing': self.point_spacing,
        #     'show_lines': self.show_lines,
        #
        #     'recent_files': self.recent_files,
        #     'subcase': subcase,
        #     # 'modes': modes,
        #     'selected_modes': selected_modes,
        #     'x_plot_type': self.x_plot_type,
        #     'plot_type': self.plot_type,
        #     'index_lim': index_lim,
        #     'eas_lim': eas_lim,
        #     'tas_lim': tas_lim,
        #     'mach_lim': mach_lim,
        #     'alt_lim': alt_lim,
        #     'q_lim': q_lim,
        #     'rho_lim': rho_lim,
        #     'ikfreq_lim': ikfreq_lim,
        #
        #     'damp_lim': ydamp_lim,
        #     'freq_lim': freq_lim,
        #     'kfreq_lim': kfreq_lim,
        #     'eigr_lim': eigr_lim,
        #     'eigi_lim': eigi_lim,
        #     'output_directory': output_directory,
        #     'units_in': units_in,
        #     'units_out': units_out,
        #     'freq_tol': freq_tol,
        #     'freq_tol_remove': freq_tol_remove,
        #     'mag_tol': mag_tol,
        #     'vl': vl,
        #     'vf': vf,
        #     'damping': damping,
        #     'damping_required': damping_required,
        #     'damping_required_tol': damping_required_tol,
        #     'eas_flutter_range': eas_flutter_range,
        #     'point_removal': point_removal,
        #     'mode_switch_method': self.mode_switch_method,
        # }
        #
        # buggy?
        # x_plot_type = self.x_plot_type
        # if x_plot_type == 'index':
        #     xlim = self.index_lim
        # elif x_plot_type == 'eas':
        #     xlim = self.eas_lim
        # elif x_plot_type == 'tas':
        #     xlim = self.tas_lim
        # elif x_plot_type == 'mach':
        #     xlim = self.mach_lim
        # elif x_plot_type == 'alt':
        #     xlim = self.alt_lim
        # elif x_plot_type == 'q':
        #     xlim = self.q_lim
        # elif x_plot_type == 'kfreq':
        #     xlim = self.kfreq_lim
        # elif x_plot_type == 'ikfreq':
        #     xlim = self.ikfreq_lim
        # else:  # pragma: no cover
        #     log.error(f'x_plot_type={x_plot_type!r} is not supported')
        #     # raise RuntimeError(x_plot_type)
        #     xlim = (None, None)

        print('settings')
        modes = None if len(self.selected_modes) == 0 else self.selected_modes
        settings = {
            'x_plot_type': 'eas',
            'nrigid_body_modes': 6,  # TODO: 6
            'f06_units': self._units_in,
            'out_units': self._units_out,
            'modes': modes,
            'VL_target': float(self.data['vl']),
            'VF_target': float(self.data['vf']),
            #'xlim_kfreq': str_limit_to_limit(self.kfreq_lim),
            #'ylim_damping': self.damping_lim,  # NO

            # self.index_lim = index_lim
            # self.tas_lim = tas_lim
            # self.mach_lim = mach_lim
            # self.alt_lim = alt_lim
            # self.q_lim = q_lim
            # self.rho_lim = rho_lim
            # self.ikfreq_lim = ikfreq_lim
            # self.ydamp_lim = ydamp_lim
            # self.kfreq_lim = kfreq_lim
            # self.freq_lim = freq_lim
            # self.eigi_lim = eigi_lim
            # self.eigr_lim = eigr_lim
            'ylim_damping': str_limit_to_limit(self.ydamp_lim),
            'ylim_freq': str_limit_to_limit(self.freq_lim),
            'eas_lim': str_limit_to_limit(self.eas_lim),
            'freq_tol': float(self.freq_tol),
            'freq_tol_remove': float(self.freq_tol_remove),
            'damping_required': float(self.damping_required),
            'damping_required_tol': double_or_blank(self.damping_required_tol, default=0.0),
            'damping_limit': float(self.damping),  # % damping
            'eas_flutter_range': str_limit_to_limit(self.eas_flutter_range),
            'plot_font_size': self.plot_font_size,
            'show_lines': self.show_lines,
            'show_points': self.show_points,
            'show_mode_number': self.show_mode_number,
            'show_detailed_mode_info': self.show_detailed_mode_info,
            'point_spacing': self.point_spacing,
            'use_rhoref': self.use_rhoref,
            'flutter_ncolumns': self.flutter_ncolumns,
            # 'mode_switch_method': None,
            #------------------
            'divergence_legend_loc': self.divergence_legend_loc,
            'flutter_bbox_to_anchor_x': self.flutter_bbox_to_anchor_x,
            'freq_ndigits': self.freq_ndigits,
            'freq_divergence_tol': self.freq_divergence_tol,
        }
        print('finished run')
        # return

        f06_filenames = out_table['Filename'].to_list()
        # print(f'f06_filenames = {f06_filenames}')
        if len(configs) == 0:
            configs = [os.path.splitext(os.path.basename(fname))[0]
                       for fname in f06_filenames]

        # print('make settings')
        print(f'settings = {settings}')

        configs, f06_filenames = remove_empty_rows(
            configs, f06_filenames, log)
        log.info(f'f06_filenames2 = {f06_filenames}')
        log.info(f'configs2 = {configs}')

        nfiles = len(f06_filenames)
        if nfiles == 0:
            log.error('no files found')
            return

        # Show progress bar and disable button
        self.progress_bar.setVisible(True)
        self.progress_bar.setMaximum(len(f06_filenames))
        self.progress_bar.setValue(0)
        self.run_organize_button.setEnabled(False)

        # Process files
        try:
            self.log.info(f'Processing {len(f06_filenames)} files...')
            write_flutter_docx(
                word_filename,
                f06_filenames, configs, out_table,
                self.log, settings,
                progress_callback=self.update_organize_progress,
                **settings)
            self.log.info(f'Successfully created {word_filename}')
        except Exception as e:
            self.log.error(f'Failed to create Word document: {str(e)}')
            self.log.error(traceback.format_exc())
        finally:
            # Hide progress bar and re-enable button
            self.progress_bar.setVisible(False)
            self.run_organize_button.setEnabled(True)

    def update_organize_progress(self, current: int, total: int):
        """Update progress bar"""
        self.progress_bar.setValue(current)
        QApplication.processEvents()  # Keep GUI responsive

    def on_base_f06_directory_load(self) -> None:
        is_passed, directory = get_file_edit('base_f06_directory', self.base_f06_directory_edit, self.log)
        if not is_passed:
            return
        if not os.path.isdir(directory):
            self.log.error(f'directory={directory!r} is not a directory')
            return
        filenames = get_files_of_type(directory, '.f06')

        # TODO: parse the filename for floats
        headers = ['Filename']
        # filenames2 = [os.path.relpath(pathi, directory) for pathi in natsort.natsorted(filenames)]
        filenames2 = list(natsort.natsorted(filenames))

        if len(filenames2) == 1:
            # can't do a trade study
            data_table = [filenames2]
        else:
            # remove the extension
            base_filenames = [os.path.splitext(os.path.basename(filename))[0] for filename in filenames2]
            data_table_initial = split_by_pattern(base_filenames)
            # print(f'base_filenames = {base_filenames}')
            # print(f'data_table_initial = {data_table_initial}')

            # define the headers
            data_row0 = data_table_initial[0]
            # print(f'data_row0 = {data_row0}')
            # headers = [f'Col{icol+1}' for icol in range(len(data_row0))] + ['Filename']
            headers = [f'{icol+1:d}' for icol in range(len(data_row0))] + ['Filename']
            # print(f'headers = {headers}')

            # add the filename onto the end
            data_table = [line + [filename] for line, filename in zip(data_table_initial, filenames2)]
            # print(f'data_table = {data_table}')
        self.table_widget.load_table_data(headers, data_table)

    @dontcrash
    def on_load_excel(self):
        is_passed, excel_filename = get_file_edit('excel_filename', self.excel_filename_edit, self.log)
        if not is_passed:
            return
        self.excel_dict = pd.read_excel(excel_filename, sheet_name=None)

        self.table_widget.clear()
        keys = list(self.excel_dict)
        # self.log.debug(f'df keys = {keys}')

        if len(keys):
            key0 = keys[0]
            # self.log.debug(f'key0 = {key0!r}')
            if self.tab_select_pulldown.count() > 0:
                self.tab_select_pulldown.clear()  # seems to cause crashes...
            self.tab_select_pulldown.addItems(keys)
            self.tab_edit.setText(key0)
            self.tab_select_pulldown.setEnabled(True)
            # self.on_select_excel_tab()

    def on_base_f06_directory_browse(self):
        start_path = self.base_f06_directory_edit.text()
        directory = self._on_load_directory(
            title='Select F06 Directory', start_path=start_path)
        if directory:
            self.base_f06_directory_edit.setText(directory)

    def on_load_excel_file(self) -> None:
        start_path = self.excel_filename_edit.text()
        filename = self._on_load_file(
            title='Select Excel File', start_path=start_path,
            file_filter='Excel File (*.xlsx);;All Files (*)')
        if filename:
            self.excel_filename_edit.setText(filename)

    def on_load_word_file(self) -> None:
        start_path = self.word_filename_edit.text()
        filename = self._on_load_file(
            title='Select Word File', start_path=start_path,
            file_filter='Word File (*.docx);;All Files (*)')
        if filename:
            self.word_filename_edit.setText(filename)

    def _on_load_file(self,
                      title: str="Select File",
                      start_path: str="",
                      file_filter: str="Text Files (*.txt);;All Files (*)") -> str:
        """
        Open a dialog to select a directory

        Parameters
        ----------
        title : str
            Dialog window title
        start_path : str
            Initial directory path

        Returns
        -------
        file_path : str
            Selected directory path, or empty string if canceled
        """
        file_path, _ = QFileDialog.getOpenFileName(self, title, start_path, file_filter)
        if not file_path:
            return ''
        return file_path

    def _on_load_directory(self,
                           title: str="Select Directory",
                           start_path: str="") -> str:
        """
        Open a dialog to select a directory

        Parameters
        ----------
        title : str
            Dialog window title
        start_path : str
            Initial directory path

        Returns
        -------
        directory : str
            Selected directory path, or empty string if canceled
        """
        directory = QFileDialog.getExistingDirectory(
            self, title, start_path,
            QFileDialog.Option.ShowDirsOnly)
        if directory:
            # print(f"Selected directory: {directory}")
            return directory
        else:
            # print("No directory selected")
            return ""

    def on_select_excel_tab(self) -> None:
        if self.tab_select_pulldown.count() == 0:
            return
        key = self.tab_select_pulldown.currentText()
        df = self.excel_dict[key]
        headers = df.columns
        data_table = df.to_numpy()
        self.table_widget.load_table_data(headers, data_table)

    def _setup_file_layout(self) -> None:
        ifile = 0
        file_row = 0
        hbox = QGridLayout()
        hbox.addWidget(self.f06_filename_label[ifile], file_row, 0)
        hbox.addWidget(self.f06_filename_edit[ifile], file_row, 1)
        hbox.addWidget(self.f06_filename_browse[ifile], file_row, 2)
        file_row += 1

        grid = QGridLayout()
        irow = 0
        grid.addWidget(self.units_in_label, irow, 0)
        grid.addWidget(self.units_in_pulldown, irow, 1)
        grid.addWidget(self.use_rhoref_checkbox, irow, 2)
        irow += 1

        grid.addWidget(self.units_out_label, irow, 0)
        grid.addWidget(self.units_out_pulldown, irow, 1)
        irow += 1

        grid.addWidget(self.subcase_label, irow, 0)
        grid.addWidget(self.subcase_edit, irow, 1)
        irow += 1

        grid.addWidget(self.x_plot_type_label, irow, 0)
        grid.addWidget(self.x_plot_type_pulldown, irow, 1)
        irow += 1
        grid.addWidget(self.plot_type_label, irow, 0)
        grid.addWidget(self.plot_type_pulldown, irow, 1)
        irow += 1

        # --------------------------------------------------
        # x-axis
        grid.addWidget(self.index_lim_label, irow, 0)
        grid.addWidget(self.index_lim_edit_min, irow, 1)
        grid.addWidget(self.index_lim_edit_max, irow, 2)
        irow += 1

        grid.addWidget(self.eas_lim_label, irow, 0)
        grid.addWidget(self.eas_lim_edit_min, irow, 1)
        grid.addWidget(self.eas_lim_edit_max, irow, 2)
        irow += 1

        grid.addWidget(self.tas_lim_label, irow, 0)
        grid.addWidget(self.tas_lim_edit_min, irow, 1)
        grid.addWidget(self.tas_lim_edit_max, irow, 2)
        irow += 1

        grid.addWidget(self.mach_lim_label, irow, 0)
        grid.addWidget(self.mach_lim_edit_min, irow, 1)
        grid.addWidget(self.mach_lim_edit_max, irow, 2)
        irow += 1

        grid.addWidget(self.alt_lim_label, irow, 0)
        grid.addWidget(self.alt_lim_edit_min, irow, 1)
        grid.addWidget(self.alt_lim_edit_max, irow, 2)
        irow += 1

        grid.addWidget(self.q_lim_label, irow, 0)
        grid.addWidget(self.q_lim_edit_min, irow, 1)
        grid.addWidget(self.q_lim_edit_max, irow, 2)
        irow += 1

        grid.addWidget(self.rho_lim_label, irow, 0)
        grid.addWidget(self.rho_lim_edit_min, irow, 1)
        grid.addWidget(self.rho_lim_edit_max, irow, 2)
        irow += 1

        grid.addWidget(self.kfreq_lim_label, irow, 0)
        grid.addWidget(self.kfreq_lim_edit_min, irow, 1)
        grid.addWidget(self.kfreq_lim_edit_max, irow, 2)
        irow += 1

        grid.addWidget(self.ikfreq_lim_label, irow, 0)
        grid.addWidget(self.ikfreq_lim_edit_min, irow, 1)
        grid.addWidget(self.ikfreq_lim_edit_max, irow, 2)
        irow += 1
        # --------------------------------------------------
        # y-axes
        grid.addWidget(self.damp_lim_label, irow, 0)
        grid.addWidget(self.damp_lim_edit_min, irow, 1)
        grid.addWidget(self.damp_lim_edit_max, irow, 2)
        irow += 1

        grid.addWidget(self.freq_lim_label, irow, 0)
        grid.addWidget(self.freq_lim_edit_min, irow, 1)
        grid.addWidget(self.freq_lim_edit_max, irow, 2)
        irow += 1

        # --------------------------------------------------
        grid.addWidget(self.eigr_lim_label, irow, 0)
        grid.addWidget(self.eigr_lim_edit_min, irow, 1)
        grid.addWidget(self.eigr_lim_edit_max, irow, 2)
        irow += 1

        grid.addWidget(self.eigi_lim_label, irow, 0)
        grid.addWidget(self.eigi_lim_edit_min, irow, 1)
        grid.addWidget(self.eigi_lim_edit_max, irow, 2)
        irow += 1
        # ------------------------------------------
        grid.addWidget(self.freq_tol_label, irow, 0)
        grid.addWidget(self.freq_tol_edit, irow, 1)
        irow += 1
        grid.addWidget(self.freq_tol_remove_label, irow, 0)
        grid.addWidget(self.freq_tol_remove_edit, irow, 1)
        irow += 1
        grid.addWidget(self.mode_label, irow, 0)
        grid.addWidget(self.mode_edit, irow, 1)
        irow += 1
        grid.addWidget(self.velocity_label, irow, 0)
        grid.addWidget(self.velocity_edit, irow, 1)
        irow += 1
        grid.addWidget(self.mag_tol_label, irow, 0)
        grid.addWidget(self.mag_tol_edit, irow, 1)
        irow += 1

        grid.addWidget(self.output_directory_label, irow, 0)
        grid.addWidget(self.output_directory_edit, irow, 1)
        grid.addWidget(self.output_directory_browse, irow, 2)
        self.output_directory_label.setDisabled(True)
        self.output_directory_edit.setDisabled(True)
        self.output_directory_browse.setDisabled(True)

        self.output_directory_label.setVisible(False)
        self.output_directory_edit.setVisible(False)
        self.output_directory_browse.setVisible(False)
        irow += 1

        grid.addWidget(self.VL_label, irow, 0)
        grid.addWidget(self.VL_edit, irow, 1)
        irow += 1

        grid.addWidget(self.VF_label, irow, 0)
        grid.addWidget(self.VF_edit, irow, 1)
        irow += 1

        grid.addWidget(self.damping_required_label, irow, 0)
        grid.addWidget(self.damping_required_edit, irow, 1)
        irow += 1

        grid.addWidget(self.damping_required_tol_label, irow, 0)
        grid.addWidget(self.damping_required_tol_edit, irow, 1)
        irow += 1

        grid.addWidget(self.damping_label, irow, 0)
        grid.addWidget(self.damping_edit, irow, 1)
        irow += 1

        grid.addWidget(self.eas_flutter_range_label, irow, 0)
        grid.addWidget(self.eas_flutter_range_edit_min, irow, 1)
        grid.addWidget(self.eas_flutter_range_edit_max, irow, 2)
        irow += 1

        # grid.addWidget(self.eas_diverg_range_label, irow, 0)
        # grid.addWidget(self.eas_diverg_range_edit_min, irow, 1)
        # grid.addWidget(self.eas_diverg_range_edit_max, irow, 2)
        # irow += 1

        grid.addWidget(self.point_removal_label, irow, 0)
        grid.addWidget(self.point_removal_edit, irow, 1)
        irow += 1

        grid.addWidget(self.mode_switch_method_label, irow, 0)
        grid.addWidget(self.mode_switch_method_pulldown, irow, 1)
        irow += 1

        jrow = 0
        grid_check = QGridLayout()
        grid_check.addWidget(self.log_xscale_checkbox, jrow, 0)
        grid_check.addWidget(self.log_yscale1_checkbox, jrow, 1)
        grid_check.addWidget(self.log_yscale2_checkbox, jrow, 2)
        jrow += 1

        grid_check.addWidget(self.show_points_checkbox[ifile], jrow, 0)
        grid_check.addWidget(self.show_mode_number_checkbox[ifile], jrow, 1)
        grid_check.addWidget(self.show_detailed_mode_info_checkbox[ifile], jrow, 2)
        jrow += 1
        grid_check.addWidget(self.point_spacing_label[ifile], jrow, 0)
        grid_check.addWidget(self.point_spacing_spinner[ifile], jrow, 1)
        jrow += 1
        grid_check.addWidget(self.include_rigid_body_modes_checkbox[ifile], jrow, 0)
        grid_check.addWidget(self.number_rigid_body_modes_label[ifile], jrow, 1)
        grid_check.addWidget(self.number_rigid_body_modes_spinner[ifile], jrow, 1)
        jrow += 1
        grid_check.addWidget(self.show_lines_checkbox[ifile], jrow, 0)
        jrow += 1

        ok_cancel_hbox = QHBoxLayout()
        ok_cancel_hbox.addWidget(self.f06_load_button)
        ok_cancel_hbox.addWidget(self.run_button)

        hbox_check = QHBoxLayout()
        hbox_check.addLayout(grid_check)
        hbox_check.addStretch(1)

        grid_modes = self._grid_modes()

        vbox = QVBoxLayout()
        vbox.addLayout(hbox)
        vbox.addLayout(grid)
        vbox.addLayout(hbox_check)
        vbox.addStretch(1)
        vbox.addLayout(ok_cancel_hbox)
        vbox.addWidget(self.pop_vtk_gui_button)
        vbox.addLayout(grid_modes)
        # log_widget = ApplicationLogWidget(self)

        log_widget = self.setup_logging()
        if self.use_dock_widgets:
            self.modes_dock_widget = NamedDockWidget('Modes', self.modes_widget, self)
            self.addDockWidget(QtCore.Qt.LeftDockWidgetArea, self.modes_dock_widget)
            self.addDockWidget(QtCore.Qt.BottomDockWidgetArea, self.log_dock_widget)
            vbox2 = vbox
        else:
            # self.log_dock_widget.hide()
            vbox2 = QHBoxLayout()
            vbox2.addWidget(self.modes_widget)
            vbox2.addLayout(vbox)
        return vbox2

    @property
    def ifile(self) -> int:
        """gets the active file index"""
        if not USE_TABS:
            return -1
        index = self.tabs.currentIndex()
        out = index - 1
        return out

    def on_new_tab(self):
        """"
        [compare, file1, +]
        """
        index = self.tabs.currentIndex()
        ifile = len(self.iwindows) - 1
        if index != ifile:
            return
        tab_plus = QWidget(self)
        iwindow_file2 = self.iwindows[-1]
        iwindow = self.tabs.addTab(tab_plus, '+')
        self.iwindows.append(iwindow)
        self.tabs.setTabText(iwindow_file2, f'File {ifile}')

    def _grid_modes(self) -> QGridLayout:
        irow = 0
        grid_modes = QGridLayout()
        grid_modes.addWidget(self.solution_type_label, irow, 0)
        grid_modes.addWidget(self.solution_type_pulldown, irow, 1)
        self.solution_type_pulldown.addItems(['Real Modes', 'Complex Modes'])
        irow += 1

        grid_modes.addWidget(self.mode2_label, irow, 0)
        grid_modes.addWidget(self.mode2_pulldown, irow, 1)
        # self.solution_type_pulldown.addItems(['Real Modes', 'Complex Modes'])
        irow += 1
        return grid_modes

    def setup_connections(self) -> None:
        self.f06_load_button.clicked.connect(self.on_load_f06)
        # self.bdf_load_button.clicked.connect(self.on_load_bdf)
        # self.op2_load_button.clicked.connect(self.on_load_op2)

        self.x_plot_type_pulldown.currentIndexChanged.connect(self.on_plot_type)
        self.plot_type_pulldown.currentIndexChanged.connect(self.on_plot_type)
        self.subcase_edit.currentIndexChanged.connect(self.on_subcase)
        self.bdf_filename_checkbox.stateChanged.connect(self.on_enable_bdf)
        self.op2_filename_checkbox.stateChanged.connect(self.on_enable_op2)
        for f06_filename_browse in self.f06_filename_browse:
            f06_filename_browse.clicked.connect(self.on_browse_f06)
        # self.modes_widget.itemSelectionChanged.connect(self.on_modes)
        # self.modes_widget.itemClicked.connect(self.on_modes)
        # self.modes_widget.currentRowChanged.connect(self.on_modes)
        self.run_button.clicked.connect(self.on_run)
        self.units_out_pulldown.currentIndexChanged.connect(self.on_units_out)
        # for ifile, box in include_rigid_body_modes_checkbox.items():
        # self.include_rigid_body_modes_checkbox[ifile].clicked.connect(self.on_rigid_body_modes)

    def on_units_out(self):
        units_out = self.units_out_pulldown.currentText()
        units_out_dict = get_flutter_units(units_out)

        tas_units = units_out_dict['velocity']
        eas_units = units_out_dict['eas']
        alt_units = units_out_dict['altitude']
        q_units = units_out_dict['dynamic_pressure']
        rho_units = units_out_dict['density']

        self.tas_lim_label.setText(f'TAS Limits ({tas_units}):')
        self.eas_lim_label.setText(f'EAS Limits ({eas_units}):')
        self.alt_lim_label.setText(f'Alt Limits ({alt_units}):')
        self.rho_lim_label.setText(f'Rho Limits ({rho_units}):')
        self.q_lim_label.setText(f'Q Limits ({q_units}):')
        self.VL_label.setText(f'VL, Limit ({eas_units}):')
        self.VF_label.setText(f'VF, Flutter ({eas_units}):')

    def on_font_size(self) -> None:
        # font_size = self.font_size_edit.value()
        self.on_set_font_size(self.font_size)

    def on_set_font_size(self, font_size: int) -> None:
        self.font_size = font_size
        font = make_font(font_size, is_bold=False)
        self.setFont(font)

    @dont_crash
    def on_load_f06(self, event) -> None:
        ifile = self.ifile
        f06_filename = os.path.abspath(self.f06_filename_edit[ifile].text())
        if not os.path.exists(f06_filename) or not os.path.isfile(f06_filename):
            self.f06_filename_edit[ifile].setStyleSheet(QLINEEDIT_RED)
            self.log.error(f"can't find {f06_filename}")
            return
        self.f06_filename_edit[ifile].setStyleSheet(QLINEEDIT_WHITE)
        f06_units = self.units_in_pulldown.currentText()
        out_units = self.units_out_pulldown.currentText()

        self.use_rhoref = self.use_rhoref_checkbox.isChecked()
        model, self.responses = load_f06_op2(
            f06_filename, self.log,
            f06_units, out_units,
            self.use_rhoref, stop_on_failure=False)

        subcases = list(self.responses.keys())
        if len(subcases) == 0:
            self.log.error('No subcases found')
            return
        # self.log.info(f'on_load_f06: subcases={subcases}')
        self.f06_filename = f06_filename
        self._units_in = f06_units
        self._units_out = out_units
        self.add_recent_file(f06_filename)
        self.update_subcases(subcases)
        self.run_button.setEnabled(True)

    def add_recent_file(self, f06_filename: str) -> None:
        path = os.path.abspath(f06_filename)
        if path in self.recent_files:
            self.recent_files.remove(path)
        self.recent_files = [path] + self.recent_files
        if len(self.recent_files) > self.nrecent_files_max:
            self.recent_files = self.recent_files[:self.nrecent_files_max]

        self._update_recent_files_actions()
        is_valid = self.validate()
        if is_valid:
            self.log.warning('add_recent_file; save')
            self._save(self.save_filename)

    def _update_recent_files_actions(self) -> None:
        current_directory = os.getcwd()
        for ifile in range(self.nrecent_files_max):
            name = f'file_{ifile}'
            try:
                fname = self.recent_files[ifile]
                show = True
            except IndexError:
                show = False

            action: QAction = self.qactions[name]
            active_text = action.text()

            if show != action.isVisible():
                action.setVisible(show)

            if show and active_text != fname:
                # only update text if it's visible
                relative_fname = os.path.relpath(fname, current_directory)
                text = fname
                if len(relative_fname) < len(fname):
                    text = relative_fname
                action.setText(text)

    def on_subcase(self) -> None:
        subcase, is_subcase_valid = self._get_subcase()
        if not is_subcase_valid:
            return
        response: FlutterResponse = self.responses[subcase]
        # self.log.info(f'on_subcase; response.results.shape={response.results.shape}')
        freqs = response.results[:, 0, response.ifreq].ravel()
        self._update_modal_participation_velocity(response)
        # self.log.info(f'on_subcase; freqs={freqs}')
        self.update_modes_table(response.modes, freqs)

    def _update_modal_participation_velocity(self, response: FlutterResponse) -> None:
        nvelocity = len(response.eigr_eigi_velocity)
        self.velocity_edit.clear()
        if nvelocity:
            velocity = response.eigr_eigi_velocity[:, -1]
            if np.all(np.isfinite(velocity)):
                items = [f'V{ipoint+1}={vel}' for ipoint, vel in enumerate(velocity)]
            else:
                items = [f'V{ipoint+1}' for ipoint in range(nvelocity)]
            self.velocity_edit.addItems(items)

    def _get_subcase(self) -> tuple[int, bool]:
        subcase_str = self.subcase_edit.currentText()
        # self.log.info(f'_get_subcase: subcase_str={subcase_str!r}')
        subcase_sline = subcase_str.split()
        # self.log.info(f'_get_subcase: subcase_sline={subcase_sline}')
        try:
            subcase = int(subcase_sline[1])
            is_valid = True
        except IndexError:
            self.log.error(f'failed parsing subcase={subcase_str!r}; subcase_sline={subcase_sline}')
            subcase = -1
            is_valid = False
        self.log.info(f'_get_subcase: subcase={subcase}; is_valid={is_valid}')
        return subcase, is_valid

    def update_subcases(self, subcases: list[int]) -> None:
        subcases_text = [f'Subcase {isubcase}' for isubcase in subcases]
        self.log.info(f'update_subcases={subcases_text}')
        self.subcase_edit.clear()
        # self.log.info(f'update_subcases setting...')
        self.subcase_edit.addItems(subcases_text)

    def update_modes_table(self, modes: list[int],
                           freqs: list[float]) -> None:
        self.modes = modes
        self._set_modes_table(self.modes_widget, modes, freqs)
        self.run_button.setEnabled(True)
        self.log.info(f'modes = {self.modes}')

    def on_modes(self) -> None:
        self.on_run()
        # self.validate()
        # self.plot(self.modes)

    # @dontcrash
    def _set_modes_table(self, modes_widget: QListWidget,
                         modes: list[int], freqs: list[float]) -> None:
        modes_widget.clear()
        for imode, freq in zip(modes, freqs):
            mode = QListWidgetItem(f'Mode {imode}; f={freq:.2f}')
            # mode.itemClicked.connect(self._on_update_mode)
            mode.setSelected(True)
            modes_widget.addItem(mode)

    def _on_update_mode(self) -> None:
        if not self.is_valid:
            # self.log.warning('_on_update_mode')
            self.validate()
        self.plot()

    def on_run(self) -> None:
        # self.log.warning('on_run')
        is_valid = self.validate()
        if not is_valid:
            return

        modes = self.selected_modes
        if len(modes) == 0:
            self.log.warning(f'modes = {modes}; assuming all modes -> {self.modes}')
            modes = self.modes
            # return
        self.log.info(f'is_valid = {is_valid}\n')
        self.is_valid = True
        self.plot(modes)
        # self.log.warning('on_run; _save')
        self._save(self.save_filename)

    @dont_crash
    def plot(self, modes: list[int]) -> None:
        log = self.log
        log.info(f'plot; modes = {modes}\n')
        if not self.is_valid:
            log.warning('not valid\n')
            return
        if len(self.responses) == 0:
            log.warning('no subcases\n')
            return

        x_plot_type = self.x_plot_type
        plot_type = self.plot_type
        log.info(f'plot_type = {plot_type}\n')

        freq_tol = self.freq_tol
        freq_tol_remove = self.freq_tol_remove
        mag_tol = self.mag_tol
        log.info(f'freq_tol = {freq_tol}\n')
        noline = not self.show_lines
        nopoints = not self.show_points
        if noline and nopoints:
            noline = False
            nopoints = True

        if x_plot_type == 'index':
            xlim = self.index_lim
        elif x_plot_type == 'eas':
            xlim = self.eas_lim
        elif x_plot_type == 'tas':
            xlim = self.tas_lim
        elif x_plot_type == 'mach':
            xlim = self.mach_lim
        elif x_plot_type == 'alt':
            xlim = self.alt_lim
        elif x_plot_type == 'q':
            xlim = self.q_lim
        elif x_plot_type == 'kfreq':
            xlim = self.kfreq_lim
        elif x_plot_type == 'ikfreq':
            xlim = self.ikfreq_lim
        else:  # pragma: no cover
            log.error(f'x_plot_type={x_plot_type!r} is not supported')
            # raise RuntimeError(x_plot_type)
            xlim = (None, None)

        # log.info(f'xlim={xlim}\n')
        if plot_type == 'zimmerman':
            print('skipping xlim check')
        else:
            assert xlim[0] != '' and xlim[1] != '', (xlim, x_plot_type)

        v_lines = get_vlines(self.vf, self.vl)
        # log.info(f'v_lines={v_lines}\n')
        # log.info(f'kfreq_lim={self.kfreq_lim}\n')
        # log.info(f'ydamp_lim={self.ydamp_lim}\n')
        # log.info(f'freq_lim={self.freq_lim}\n')
        # log.info(f'damping={self.damping}\n')
        xlim_kfreq = self.kfreq_lim
        ylim_damping = self.ydamp_lim
        ylim_freq = self.freq_lim

        damping_required = self.damping_required
        damping_required_tol = self.damping_required_tol
        damping_limit = self.damping  # % damping
        eas_flutter_range = self.eas_flutter_range
        # eas_diverg_range = self.eas_diverg_range
        if damping_required_tol is None:
            damping_required_tol = 0.01
        if damping_required_tol < 0.0:
            damping_required_tol = 0.0

        # changing directory so we don't make a long filename
        # in the plot header
        # log.info(f'damping_limit = {damping_limit}\n')
        dirname = os.path.abspath(os.path.dirname(self.f06_filename))
        basename = os.path.basename(self.f06_filename)

        base = os.path.splitext(basename)[0]
        current_directory = os.getcwd()
        sys.stdout.flush()
        os.chdir(dirname)

        fig = plt.figure(1)
        fig.clear()
        log.info(f'cleared plot\n')
        if plot_type not in {'root-locus', 'modal-participation', 'zimmerman'}:
            gridspeci = gridspec.GridSpec(2, 4)
            damp_axes = fig.add_subplot(gridspeci[0, :3])
            freq_axes = fig.add_subplot(gridspeci[1, :3], sharex=damp_axes)

        response = self.responses[self.subcase]

        # you can change the output units without reloading
        if self._units_out != self.units_out:
            response.convert_units(self.units_out)
            self._units_out = self.units_out

        response.noline = noline
        response.freq_ndigits = self.freq_ndigits
        response.set_symbol_settings(
            nopoints, self.show_mode_number, self.point_spacing)
        # log.info(f'self.plot_font_size = {self.plot_font_size}')
        response.set_font_settings(self.plot_font_size)
        response.log = log
        # print('trying plots...')

        # log.info(f'getting logs\n')
        log_scale_x = self.data['log_scale_x']
        log_scale_y1 = self.data['log_scale_y1']
        log_scale_y2 = self.data['log_scale_y2']
        # print(f'log_scale_x={log_scale_x}; log_scale_y1={log_scale_y1}; log_scale_y2={log_scale_y2}')
        # print(f'export_to_png={self.export_to_png}')

        # print(f'point_removal = {self.point_removal}')
        png_filename0, png_filename = get_png_filename(
            base, x_plot_type, plot_type,
            self.export_to_png)
        print(f'png_filename={png_filename}')
        try:
            if plot_type == 'zimmerman':
                axes = fig.add_subplot(111)
                response.plot_zimmerman(fig=fig, axes=axes, modes=modes, show=True)
            elif plot_type == 'root-locus':
                axes = fig.add_subplot(111)
                # log.info(f'modes={modes}; eigr_lim={self.eigr_lim}; eigi_lim={self.eigi_lim}; freq_tol={freq_tol}')
                # log.info(f'png_filename={png_filename}')
                response.plot_root_locus(
                    fig=fig, axes=axes,
                    modes=modes, eigr_lim=self.eigr_lim, eigi_lim=self.eigi_lim,
                    freq_tol=freq_tol, freq_tol_remove=freq_tol_remove,
                    show=True, clear=False, close=False,
                    legend=True,
                    png_filename=png_filename,
                )
            elif plot_type == 'modal-participation':
                axes = fig.add_subplot(111)
                mode = self.mode_edit.value()
                ivel = self.velocity_edit.currentIndex()
                # print(f'ivel={ivel}; mode={mode}')
                response.plot_modal_participation(
                    ivel, mode,
                    fig=fig, axes=axes,
                    modes=modes,  # eigr_lim=self.eigr_lim, eigi_lim=self.eigi_lim,
                    freq_tol=freq_tol, freq_tol_remove=freq_tol_remove,
                    mag_tol=mag_tol,
                    show=True, clear=False, close=False,
                    legend=True,
                    png_filename=png_filename,
                )
            elif plot_type == 'x-damp-kfreq':
                # xlabel: eas
                # ylabel1 = r'Structural Damping; $g = 2 \gamma $'
                # ylabel2 = r'KFreq [rad]; $ \omega c / (2 V)$'
                # print('plot_kfreq_damping')
                response.plot_kfreq_damping(
                    fig=fig, damp_axes=damp_axes, freq_axes=freq_axes,
                    modes=modes, plot_type=x_plot_type,
                    xlim=xlim, ylim_damping=ylim_damping, ylim_kfreq=xlim_kfreq,
                    freq_tol=freq_tol, freq_tol_remove=freq_tol_remove,
                    show=True, clear=False, close=False,
                    legend=True,
                    png_filename=png_filename,
                )
                update_ylog_style(fig, log_scale_x, log_scale_y1, log_scale_y2)
            else:
                assert plot_type in 'x-damp-freq', plot_type
                # print('plot_vg_vf')
                # log.info(f'png_filename={png_filename!r}')
                # log.info(f'modes={modes!r}')
                # log.info(f'freq_tol={freq_tol!r}')
                # log.info(f'v_lines={v_lines!r}')
                damping_crossings = get_damping_crossings(
                    damping_required, damping_required_tol, damping_limit)

                response.plot_vg_vf(
                    fig=fig, damp_axes=damp_axes, freq_axes=freq_axes,
                    plot_type=x_plot_type,
                    modes=modes,
                    xlim=xlim, ylim_damping=ylim_damping, ylim_freq=ylim_freq,
                    eas_range=eas_flutter_range,
                    freq_tol=freq_tol, freq_tol_remove=freq_tol_remove,
                    show=True, clear=False, close=False,
                    legend=True,
                    v_lines=v_lines,
                    damping_limit=damping_limit,
                    damping_required=damping_required,
                    damping_crossings=damping_crossings,
                    png_filename=png_filename,
                    point_removal=self.point_removal,
                    mode_switch_method=self.mode_switch_method,
                    show_detailed_mode_info=self.show_detailed_mode_info,
                    ncol=self.flutter_ncolumns,
                    divergence_legend_loc=self.divergence_legend_loc,
                    flutter_bbox_to_anchor=(self.flutter_bbox_to_anchor_x, 1.),
                )
                update_ylog_style(fig, log_scale_x, log_scale_y1, log_scale_y2)
                fig.canvas.draw()
        except Exception as e:  # pragma: no cover
            log.error(f'plot_type={plot_type}')
            log.error(str(e))
            print(traceback.format_exc())
            # print(traceback.print_tb())
            print(traceback.print_exception(e))
            raise

        export_flutter_results(
            response, modes, png_filename0,
            self.export_to_csv, self.export_to_zaero, self.export_to_f06,
            log)
        os.chdir(current_directory)
        if png_filename:
            log.info(f'saved {png_filename}')
        else:
            log.info(f'did not write file because export_to_png=False')

    def get_xlim(self) -> tuple[Limit, Limit, Limit, Limit,
                                Limit, Limit, Limit, Limit, Limit,
                                Optional[float], Optional[float], Optional[float],
                                Optional[float], Optional[float], bool]:
        index_lim, is_passed0 = get_list_float_or_none(
            [self.index_lim_edit_min, self.index_lim_edit_max])
        eas_lim, is_passed1 = get_list_float_or_none(
            [self.eas_lim_edit_min, self.eas_lim_edit_max])
        tas_lim, is_passed2 = get_list_float_or_none(
            [self.tas_lim_edit_min, self.tas_lim_edit_max])

        mach_lim, is_passed3 = get_list_float_or_none(
            [self.mach_lim_edit_min, self.mach_lim_edit_max])
        alt_lim, is_passed4 = get_list_float_or_none(
            [self.alt_lim_edit_min, self.alt_lim_edit_max])
        q_lim, is_passed5 = get_list_float_or_none(
            [self.q_lim_edit_min, self.q_lim_edit_max])
        rho_lim, is_passed6 = get_list_float_or_none(
            [self.rho_lim_edit_min, self.rho_lim_edit_max])

        is_passed_x = all([
            is_passed0, is_passed1, is_passed2, is_passed3,
            is_passed4, is_passed5, is_passed6,
        ])
        damp_lim, is_passed_damp = get_list_float_or_none(
            [self.damp_lim_edit_min, self.damp_lim_edit_max])

        freq_lim, is_passed_freq = get_list_float_or_none(
            [self.freq_lim_edit_min, self.freq_lim_edit_max])
        kfreq_lim, is_passed_kfreq = get_list_float_or_none(
            [self.kfreq_lim_edit_min, self.kfreq_lim_edit_max])
        ikfreq_lim, is_passed_ikfreq = get_list_float_or_none(
            [self.ikfreq_lim_edit_min, self.ikfreq_lim_edit_max])
        eigr_lim, is_passed_eigr = get_list_float_or_none(
            [self.eigr_lim_edit_min, self.eigr_lim_edit_max])
        eigi_lim, is_passed_eigi = get_list_float_or_none(
            [self.eigr_lim_edit_min, self.eigr_lim_edit_max])
        is_passed_eig = all([is_passed_eigr, is_passed_eigi])

        freq_tol, is_passed_tol1 = get_float_or_none(self.freq_tol_edit)
        freq_tol_remove, is_passed_tol2 = get_float_or_none(self.freq_tol_remove_edit)
        mag_tol, is_passed_tol3 = get_float_or_none(self.mag_tol_edit)
        if is_passed_tol1 and freq_tol is None:
            freq_tol = -1.0
        if is_passed_tol2 and freq_tol_remove is None:
            freq_tol_remove = -1.0
        if is_passed_tol3 and mag_tol is None:
            mag_tol = -1.0

        vl, is_passed_vl = get_float_or_none(self.VL_edit)
        vf, is_passed_vf = get_float_or_none(self.VF_edit)
        damping_required, is_passed_damping_required = get_float_or_none(self.damping_required_edit)
        damping_required_tol, is_passed_damping_required_tol = get_float_or_none(self.damping_required_tol_edit)
        damping, is_passed_damping = get_float_or_none(self.damping_edit)
        eas_flutter_range, is_passed_flutter_range = get_list_float_or_none(
            [self.eas_flutter_range_edit_min, self.eas_flutter_range_edit_max])

        # 595:596,600:601
        point_removal_str = self.point_removal_edit.text().strip()
        point_removal = point_removal_str_to_point_removal(point_removal_str, self.log)

        vl = _float_passed_to_default(vl, is_passed_vl)
        vf = _float_passed_to_default(vf, is_passed_vf)
        log = self.log
        damping_required = _float_passed_to_default(damping_required, is_passed_damping_required)
        damping = _float_passed_to_default(damping, is_passed_damping)

        # if is_passed_eas_damping_lim1 and eas_damping_lim_min is None:
        #    eas_damping_lim_min = None
        # if is_passed_eas_damping_lim2 and eas_damping_lim_max is None:
        #    eas_damping_lim_max = None

        # self.log.info(f'XLim = {xlim}')
        is_passed_flags = [
            is_passed_x,
            is_passed_damp,
            is_passed_freq, is_passed_kfreq, is_passed_ikfreq,
            is_passed_eig,
            is_passed_tol1, is_passed_tol2, is_passed_tol3,
            is_passed_vl, is_passed_vf,
            is_passed_damping_required, is_passed_damping_required_tol,
            is_passed_damping,
            is_passed_flutter_range,
        ]
        is_passed = all(is_passed_flags)
        # if not is_passed:
        # self.log.warning(f'is_passed_flags = {is_passed_flags}')
        # print(f'freq_tol = {freq_tol}')
        out = (
            index_lim, eas_lim, tas_lim, mach_lim, alt_lim, q_lim, rho_lim,
            damp_lim, freq_lim, kfreq_lim, ikfreq_lim,
            eigr_lim, eigi_lim,
            freq_tol, freq_tol_remove, mag_tol,
            vl, vf, damping_required, damping_required_tol, damping,
            eas_flutter_range, point_removal, is_passed,
        )
        return out

    def get_selected_modes(self) -> list[int]:
        mode_strs = get_selected_items_flat(self.modes_widget)
        # self.log.info(f'mode_strs = {mode_strs}')
        modes = [int(mode_str.split(';')[0].split(' ')[1])
                 for mode_str in mode_strs]
        self.log.info(f'modes = {modes}')
        return modes

    def validate(self) -> bool:
        ifile = self.ifile
        # self.log.warning('validate')
        (index_lim, eas_lim, tas_lim, mach_lim, alt_lim, q_lim, rho_lim,
         ydamp_lim, freq_lim, kfreq_lim, ikfreq_lim,
         eigr_lim, eigi_lim,
         freq_tol, freq_tol_remove, mag_tol,
         vl, vf, damping_required, damping_required_tol, damping,
         eas_flutter_range, point_removal,
         is_valid_xlim) = self.get_xlim()

        selected_modes = []
        subcase, is_subcase_valid = self._get_subcase()
        self.log.warning(f'subcase={subcase}; is_subcase_valid={is_subcase_valid}')
        if is_subcase_valid:
            selected_modes = self.get_selected_modes()

        self.subcase = subcase
        self.selected_modes = selected_modes
        self.index_lim = index_lim
        self.eas_lim = eas_lim
        self.tas_lim = tas_lim
        self.mach_lim = mach_lim
        self.alt_lim = alt_lim
        self.q_lim = q_lim
        self.rho_lim = rho_lim
        self.ikfreq_lim = ikfreq_lim
        self.ydamp_lim = ydamp_lim
        self.kfreq_lim = kfreq_lim
        self.freq_lim = freq_lim
        self.eigi_lim = eigi_lim
        self.eigr_lim = eigr_lim
        self.freq_tol = freq_tol
        self.freq_tol_remove = freq_tol_remove
        self.mag_tol = mag_tol
        self.vl = vl
        self.vf = vf
        self.damping_required = damping_required
        self.damping_required_tol = damping_required_tol
        self.damping = damping
        self.eas_flutter_range = eas_flutter_range
        self.point_removal = point_removal

        self.x_plot_type = self.x_plot_type_pulldown.currentText()
        self.plot_type = self.plot_type_pulldown.currentText()
        self.mode_switch_method = self.mode_switch_method_pulldown.currentText()

        units_in = self.units_in_pulldown.currentText()
        units_out = self.units_out_pulldown.currentText()
        output_directory = self.output_directory_edit.text()

        self.show_lines = self.show_lines_checkbox[ifile].isChecked()
        self.show_points = self.show_points_checkbox[ifile].isChecked()
        self.show_mode_number = self.show_mode_number_checkbox[ifile].isChecked()
        self.show_detailed_mode_info = self.show_detailed_mode_info_checkbox[ifile].isChecked()
        self.point_spacing = self.point_spacing_spinner[ifile].value()
        self.use_rhoref = self.use_rhoref_checkbox.isChecked()

        is_passed_modal_partipation = False
        subcases = list(self.responses)
        if len(subcases):
            self.log.info(f'subcases={subcases}')
            subcase0 = subcases[0]
            response = self.responses[subcase0]

            failed_modal_partipation = (
                (self.plot_type == 'modal-participation') and
                ((response.eigr_eigi_velocity is None) or
                 (response.eigenvector is None))
            )
            is_passed_modal_partipation = not failed_modal_partipation
        # (
        #     (self.plot_type == 'modal-participation') and
        #     (response.eigr_eigi_velocity is not None)
        # ) or (self.plot_type != 'modal-participation'))
        data = {
            # 'bdf_filename': self.bdf_filename,
            # 'op2_filename': self.op2_filename,
            'log_scale_x': self.log_xscale_checkbox.isChecked(),
            'log_scale_y1': self.log_yscale1_checkbox.isChecked(),
            'log_scale_y2': self.log_yscale2_checkbox.isChecked(),
            'use_rhoref': self.use_rhoref,
            'show_points': self.show_points,
            'show_mode_number': self.show_mode_number,
            'show_detailed_mode_info': self.show_detailed_mode_info,
            'point_spacing': self.point_spacing,
            'show_lines': self.show_lines,

            'recent_files': self.recent_files,
            'subcase': subcase,
            # 'modes': modes,
            'selected_modes': selected_modes,
            'x_plot_type': self.x_plot_type,
            'plot_type': self.plot_type,
            'index_lim': index_lim,
            'eas_lim': eas_lim,
            'tas_lim': tas_lim,
            'mach_lim': mach_lim,
            'alt_lim': alt_lim,
            'q_lim': q_lim,
            'rho_lim': rho_lim,
            'ikfreq_lim': ikfreq_lim,

            'damp_lim': ydamp_lim,
            'freq_lim': freq_lim,
            'kfreq_lim': kfreq_lim,
            'eigr_lim': eigr_lim,
            'eigi_lim': eigi_lim,
            'output_directory': output_directory,
            'units_in': units_in,
            'units_out': units_out,
            'freq_tol': freq_tol,
            'freq_tol_remove': freq_tol_remove,
            'mag_tol': mag_tol,
            'vl': vl,
            'vf': vf,
            'damping': damping, # 0.03
            'damping_required': damping_required, # 0.0
            'damping_required_tol': damping_required_tol, # 0.0001
            'eas_flutter_range': eas_flutter_range,
            'point_removal': point_removal,
            'mode_switch_method': self.mode_switch_method,
        }
        self.units_in = units_in
        self.units_out = units_out
        is_passed = all([is_valid_xlim, is_subcase_valid, is_passed_modal_partipation])
        if is_passed:
            self.data = data
            # self.xlim = xlim
            # self.ylim = ydamp_lim
            # self.data = data
            # is_valid = validate_json(self.data, self.log)
            # if is_valid != is_passed:
            # self.log.info(f'passed data:\n{str(self.data)}')
        else:
            del data['recent_files']
            self.log.error(
                f'is_valid_xlim = {is_valid_xlim}\n'
                f'is_subcase_valid = {is_subcase_valid}\n'
                f'is_passed_modal_partipation = {is_passed_modal_partipation}\n'
                f'failed data:\n{str(data)}'
            )
            # self.log.error(f'failed data:\n{str(data)}')
        return is_passed

    def log_debug(self, msg: str) -> None:
        print(f'DEBUG: {msg}')

    def log_info(self, msg: str) -> None:
        print(f'INFO:  {msg}')

    def log_command(self, msg: str) -> None:
        print(f'COMMAND: {msg}')

    def log_warning(self, msg: str) -> None:
        print(f'WARNING: {msg}')

    def log_error(self, msg: str) -> None:
        print(f'ERROR:   {msg}')


def main(f06_filename: str='') -> None:  # pragma: no cover
    # kills the program when you hit Cntl+C from the command line
    # doesn't save the current state as presumably there's been an error
    import signal
    signal.signal(signal.SIGINT, signal.SIG_DFL)

    import sys
    # Someone is launching this directly
    # Create the QApplication
    app = QApplication(sys.argv)
    # The Main window

    main_window = FlutterGui(f06_filename)
    # Enter the main loop
    app.exec_()


def get_file_edit(name: str,
                  filename_edit: QLineEdit,
                  log: SimpleLogger) -> tuple[bool, str]:
    is_passed = False
    filename = filename_edit.text()
    if not os.path.exists(filename):
        log.error(print_bad_path(filename))
        return is_passed, ''
    log.info(f'loading {name} {filename}')
    is_passed = True
    return is_passed, filename


def write_flutter_docx(docx_filename: str,
                       f06_filenames: list[str],
                       configs: list[str],
                       table: pd.DataFrame,
                       log: SimpleLogger,
                       settings: dict[str, int | float | str],
                       x_plot_type: str='eas',
                       f06_units: str='english_in',
                       out_units: str='english_kt',
                       nrigid_body_modes: int=0,
                       modes=None,
                       VL_target=None,
                       VF_target=None,
                       # v_lines=None,
                       # xlim_kfreq=None,
                       ylim_damping=None,
                       ylim_freq=None,
                       eas_lim=None,
                       freq_tol=None,
                       freq_tol_remove=None,
                       damping_required=None,
                       damping_required_tol: float=0.0001,
                       damping_limit=None,
                       eas_flutter_range=None,
                       plot_font_size: int=10,
                       show_lines: bool=True,
                       show_points: bool=True,
                       show_mode_number: bool=False,
                       show_detailed_mode_info: bool=False,
                       point_spacing: int=8,
                       use_rhoref: bool=False,
                       flutter_ncolumns=None,
                       divergence_legend_loc=None,
                       flutter_bbox_to_anchor_x=None,
                       freq_ndigits=None,
                       freq_divergence_tol=None,
                       progress_callback=None,
                       ) -> None:
    f06_filename0 = f06_filenames[0]
    dirname = os.path.dirname(f06_filename0)
    docx_filename = os.path.join(dirname, docx_filename)

    #------------------------------
    mode_switch_method = None
    x_cutoff = None  # TODO: add me; fixes NaNs?
    # point_spacing = 8
    make_pngs = True
    show_individual = False
    freq_target = 0.5
    V_baseline = 1000.0
    # VL_target, VF_target
    #------------------------------
    eas_range = eas_flutter_range
    ncol = flutter_ncolumns
    #------------------------------
    v_lines = get_vlines(VF_target, VL_target)

    if ncol in [0, 1]:
        figsize = (15, 12)
    else:
        figsize = (24, 12)

    noline = not show_lines
    nopoints = not show_points
    if noline and nopoints:
        noline = False
        nopoints = True

    # damping_required = None,
    # damping_required_tol = None,

    damping_crossings, damping_required_tol = _get_damping_crossings(
        damping_required, damping_required_tol,
        damping_limit)
    settings['damping_required_tol'] = str(damping_required_tol)
    settings['damping_crossings'] = str(damping_crossings)

    #------------------------------
    # damping_crossings2 = {}
    # if isinstance(damping_crossings, list):
    #     for key, value in damping_crossings:
    #         damping_crossings2[key] = value
    #     damping_crossings = damping_crossings2
    #     damping_crossings2 = {}
    #
    # for key, value in damping_crossings.items():
    #     if np.allclose(key, 0.0):
    #         damping_crossings2[key] = key + 0.001
    #     damping_crossings2[key] = value

    #------------------------------
    cases = []

    nfiles = len(f06_filenames)
    log.info(f'f06_filenames = {f06_filenames}')
    log.info(f'configs = {configs}')
    for ifile, f06_filename, config in zip(count(), f06_filenames, configs):
        log.info(f'Processing F06 {ifile}/{nfiles}: {f06_filename}')
        if progress_callback is not None:
            progress_callback(ifile, nfiles)  # 0-indexed for progress bar

        base = os.path.splitext(f06_filename)[0]
        png_filename = base + '.png'
        if config.strip() == '':
            config = os.path.basename(base)

        resp_dict, data_dict = make_flutter_response(
            str(f06_filename),
            f06_units=f06_units, out_units=out_units,
            use_rhoref=use_rhoref,
            log=log)
        assert len(resp_dict) == 1, resp_dict
        response = resp_dict[1]

        # xcutoff doesn't apply for first 6 modes
        response.nrigid_body_modes = nrigid_body_modes
        response.x_cutoff = x_cutoff
        # response.noline = noline
        response.freq_ndigits = freq_ndigits

        response.set_plot_settings(
            figsize=figsize,
            # the delta spacing for the x/y axis
            # xtick_major_locator_multiple=[50., 50.],
            # ytick_major_locator_multiple=[0.02, None],
        )
        response.set_symbol_settings(
            nopoints=False, show_mode_number=show_mode_number,
            point_spacing=point_spacing, markersize=5,
        )
        response.set_font_settings(plot_font_size)

        vl_vf_crossing_dict, vd_crossing_dict = response.get_flutter_crossings(
            damping_crossings=damping_crossings, modes=modes,
            eas_range=eas_range)

        if make_pngs:
            log.info(f"modes in plot_vg_vf = {modes}")
            fig, (damp_axes, freq_axes) = response.plot_vg_vf(
                plot_type='eas', modes=modes,
                clear=False, close=False, legend=True,
                xlim=eas_lim, ylim_damping=ylim_damping, ylim_freq=ylim_freq,
                # ivelocity: Optional[int]=None,
                v_lines=v_lines,
                damping_required=damping_required,
                damping_limit=damping_limit,
                ncol=ncol, freq_tol=freq_tol, freq_tol_remove=freq_tol_remove,
                #--------
                mode_switch_method=mode_switch_method,
                divergence_legend_loc=divergence_legend_loc,
                flutter_bbox_to_anchor=(flutter_bbox_to_anchor_x, 1.),
                # plot_freq_tol_filtered_lines=True,
                damping_crossings=damping_crossings, filter_damping=True,
                eas_range=eas_range,
                png_filename=None,
                filter_freq=True,
                show_detailed_mode_info=show_detailed_mode_info,
                show=False)
            basename = os.path.basename(png_filename)
            damp_axes.set_title(basename)
            # plt.tight_layout()
            # if show_individual:
            #     plt.show()
            fig.savefig(png_filename, bbox_inches='tight')
            # bbox_to_anchor=(1, 1), borderaxespad=0)
            # shutil.copyfile(png_filename, png_filename_mach)
            plt.close()
        if show_individual:
            raise RuntimeError('stopping')

        # print(f'modes = {modes}')
        # print(f'xcrossing_dict = {xcrossing_dict}')
        hump_message = _get_hump_message(
            response, vl_vf_crossing_dict,
            eas_range, show_detailed_mode_info)

        log.info(f'VL_target = {VL_target}')
        v0, freq0, v3, freq3, vdiverg, freq_diverg = response.xcrossing_dict_to_VL_VF_VD(
            vl_vf_crossing_dict, vd_crossing_dict,
            log, freq_target, VL_target, VF_target,
            v_baseline=V_baseline,
            # is_hump_modes=parse_hump_modes,
        )
        # if VL < VL_target:
        #     log.error(f'VL={VL} KEAS, freq={freqL} Hz; {f06_filename_base}')
        # if VF < VF_target:
        #     log.error(f'VF={VF} KEAS, freq={freqF} Hz; {f06_filename_base}')
        # if VD < VD_target:
        #     log.error(f'VD={VD} KEAS, freq={freqD} Hz; {f06_filename_base}')
        # mass = -1.0
        # cg = [0., 0., 0.]
        # inertia = [0., 0., 0., 0., 0., 0.]

        if 'opgwg' not in data_dict:
            matrices = data_dict['matrices']
            log.warning(f'data_dict_keys={list(data_dict)}; matrices_keys={list(matrices)}')
            mass = np.full(1, np.nan)
            cg = np.full(3, np.nan)
            inertia = np.full((3, 3), np.nan)
        else:
            opgwg = data_dict['opgwg']  # grid point weight
            # matrices = data_dict['matrices']
            # frequencies = matrices['freq']
            mass = opgwg['mass']
            cg = opgwg['cg']
            # print(opgwg)
            # print(f'frequencies = {frequencies.round(3)}')
            inertia = opgwg['I(S)']

        case = (v0, freq0, v3, freq3, vdiverg, freq_diverg,
                mass, cg, inertia, config,
                hump_message,
                f06_filename, png_filename)
        eas_units = response.out_units['eas']
        cases.append(case)
    _cases_to_document(log, docx_filename, cases, settings,
                       eas_units=eas_units)


def _get_damping_crossings(damping_required: float,
                           damping_required_tol: Optional[float],
                           damping_limit: float,
                           ) -> tuple[dict[float, float], float]:
    damping_crossings = {}
    if damping_required_tol is None or damping_required_tol < 0.:
        damping_required_tol = 0.0

    if damping_required is not None and damping_required > -1.0:
        # VL
        damping_required_tol = max(0., damping_required_tol)
        damping_crossings[damping_required] = damping_required + damping_required_tol
    if damping_limit is not None and damping_limit > -1.0:
        # VF
        damping_crossings[damping_limit] = damping_limit
    return damping_crossings, damping_required_tol

def _cases_to_document(log: SimpleLogger,
                       docx_filename: PathLike,
                       cases: list, settings: dict[str, int | float],
                       eas_units: str='KEAS',
                       write_filename: bool=True):
    percent0 = settings['damping_required'] * 100
    percent3 = settings['damping_limit'] * 100
    label_vg0 = f'V,g={percent0:.0f}% ({eas_units})'
    label_vg3 = f'V,g={percent3:.0f}% ({eas_units})'
    label_vd = f'VDiverg ({eas_units})'

    label_freq_g0 = f'Freq,g={percent0:.0f}% (Hz)'
    label_freq_g3 = f'Freq,g={percent3:.0f}% (Hz)'

    configs = []
    f06_filenames = []
    # config_file_table = {
    #     'Config': configs,
    #     'File': f06_filenames,
    # }

    flutter_table = {
        # 'Configuration': [],
        'Config': configs,
        'File': f06_filenames,
        label_vg0: [],
        label_freq_g0: [],
        label_vg3: [],
        label_freq_g3: [],
        label_vd: [],
    }

    document = Document()
    _write_name_value_table(document, settings)

    for case in cases:
        # document.add_heading(f'Config={config0}', heading_level_mach)

        (v0, freq0, v3, freq3, vdiverg, freq_diverg,
         mass, cg, inertia, config,
         hump_message,
         f06_filename, png_filename) = case

        # could get the config name better
        basename = os.path.basename(f06_filename)

        configs.append(config)
        f06_filenames.append(f06_filename)
        flutter_table[label_vg0].append(v0)
        flutter_table[label_vg3].append(v3)
        flutter_table[label_vd].append(vdiverg)
        flutter_table[label_freq_g0].append(freq0)
        flutter_table[label_freq_g3].append(freq3)

        # configi = ''
        if np.isfinite(freq0):
            v0_text = f'V 0%={v0:.0f} KEAS ({freq0:.1f} Hz)'
        else:  # TODO: why does this happen?
            v0_text = f'V 0%={v0:.0f} KEAS (default)'

        if np.isfinite(freq3):
            v3_text = f'V 3%={v3:.0f} KEAS ({freq3:.1f} Hz)'
        else:  # TODO: why does this happen?
            v3_text = f'V 3%={v3:.0f} KEAS (default)'

        text = f'{v0_text}, {v3_text}, VD={vdiverg:.0f} KEAS, {config}'
        if hump_message:
            text += '\n' + hump_message

        freq0_str = f'{freq0:.2f}' if np.isfinite(freq0) else 'N/A'
        freq3_str = f'{freq3:.2f}' if np.isfinite(freq3) else 'N/A'

        # write the dirname and f06_filename of the file
        dirname = os.path.basename(os.path.dirname(f06_filename))
        path_str0 = os.path.join(dirname, basename)
        path_str = str(path_str0.replace('\\', '/'))

        if os.path.exists(png_filename):
            document.add_picture(str(png_filename), width=Inches(6.5))
        else:
            paragraph = document.add_paragraph(f'Missing {f06_filename}')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if write_filename:
            paragraph = document.add_paragraph(path_str)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # msg += f'{mach:.3f}, {weight_config}, {v0:.3f}, {freq0_str}, {v3:.3f}, {freq3_str}, {vdiverg:.3f}, {config}, {f06_filename_base}\n'

    if percent0 <= -100.0:
        del flutter_table[label_vg0]
        del flutter_table[label_freq_g0]
    if percent3 <= -100.0:
        del flutter_table[label_vg3]
        del flutter_table[label_freq_g3]
    _write_2d_table(document, flutter_table, log, 'Flutter Results')
    log.info(f'saving docx {docx_filename}')
    document.save(docx_filename)
    return
    # ncol = flutter_ncolumns
    #
    # damping_crossings = [
    #     (0.00, 0.01),
    #     (0.03, 0.03),
    # ]
    #
    # damping_crossings2 = {}
    # if isinstance(damping_crossings, list):
    #     for key, value in damping_crossings:
    #         damping_crossings2[key] = value
    #     damping_crossings = damping_crossings2
    #     damping_crossings2 = {}
    #
    # for key, value in damping_crossings.items():
    #     if np.allclose(key, 0.0):
    #         damping_crossings2[key] = key + 0.001
    #     damping_crossings2[key] = value
    #
    # if filename_base_to_vl_vf_vbase_dict is None:
    #     filename_base_to_vl_vf_vbase_dict = {}
    # assert isinstance(filename_base_to_vl_vf_vbase_dict, dict), filename_base_to_vl_vf_vbase_dict
    # print(f'filename_base_to_vl_vf_vbase_dict = {filename_base_to_vl_vf_vbase_dict}')
    #
    # if case_format == 'func':
    #     assert case_info_func is not None, case_info_func
    #
    # if x_cutoff is None:
    #     x_cutoff = eas_lim[1]
    #
    # # type checking
    # str(eas_lim[0] + 1)
    # str(eas_lim[1] + 1)
    #
    # if sub_dirnames is None:
    #     log.warning(f'sub_dirnames = {sub_dirnames}; assuming single folder')
    #     sub_dirnames = [None]
    # if skip_extra_cases:
    #     log.error(f'skip_extra_cases={skip_extra_cases}; should be temporary')
    #
    # # ---------------------------------------------------------------------------------
    # cases = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
    # cases2 = defaultdict(list)
    # for sub_dirname in sub_dirnames:
    #     sub_dirname2: str = '' if sub_dirname is None else sub_dirname
    #     # if sub_dirname is None:
    #     #     dirnamei = dirname
    #     # else:
    #     dirnamei = dirname / sub_dirname2
    #
    #     dirnamei_pics = dirname / f'pics_{sub_dirname}'
    #     print(dirnamei_pics)
    #     # if not dirnamei_pics.exists():
    #     os.makedirs(dirnamei_pics, exist_ok=True)
    #
    #     print(f'directory = {str(dirnamei)}')
    #     print(f'glob_str = {str(glob_str)}')
    #     # f06_root_filenames = [fname for fname in os.listdir(dirnamei) if fname.endswith('.f06')]
    #     f06_root_filenames = dirnamei.glob('*.f06')
    #     f06_root_filenames = natsorted(f06_root_filenames)
    #     f06_root_filenames = filter_f06(f06_root_filenames, glob_str)
    #     # f06_filenames = [dirnamei / fname for fname in f06_root_filenames]
    #     f06_filenames = f06_root_filenames
    #
    #     for f06_filename in f06_filenames:
    #         print('------------------------------------------------')
    #         assert os.path.exists(f06_filename), print_bad_path(f06_filename)
    #
    #         f06_filename_base = f06_filename.name
    #
    #         base1 = os.path.splitext(f06_filename_base)[0]
    #         png_filename = dirnamei_pics / (base1 + '.png')
    #         print(f'f06_filename = {str(f06_filename)}')
    #         # print(f'png_filename = {str(png_filename)}')
    #         is_failed, out = get_case_info(
    #             f06_filename_base,
    #             store_configs,
    #             weight_configs,
    #             machs_to_parse, machs_to_skip,
    #             case_format, case_info_func, skip_extra_cases)
    #         if is_failed:
    #             continue
    #         mach_num_str = out['mach_num_str']
    #         mach_number = out['mach']
    #         weight_config = out['weight_config']
    #         configi = out['config']
    #         # mach_num_str, mach_number, weight_config, configi = out
    #
    #         log.level = 'debug'
    #         # print(f'f06_filename_base = {f06_filenme_base}')
    #         try:
    #             resp_dict, data_dict = make_flutter_response(
    #                 str(f06_filename),
    #                 f06_units='english_in', out_units='english_kt',
    #                 use_rhoref=use_rhoref,
    #                 log=log)
    #             assert len(resp_dict) == 1, resp_dict
    #         except:
    #             raise
    #             if stop_on_failure:
    #                 raise
    #             VL = -1.0
    #             VF = -1.0
    #             VD = -1.0
    #             freqL = -1.0
    #             freqF = -1.0
    #             # freqD = -1.0
    #             mass = np.zeros(1)[0]
    #             cg = np.zeros(3)
    #             inertia = np.zeros((3, 3))
    #             png_filename = ''
    #             case_key = (mach_number, weight_config, store_config)
    #             case_value = (VL, freqL, VF, freqF, VD,
    #                           mass, cg, inertia, configi,
    #                           '',
    #                           f06_filename_base, f06_filename, png_filename)
    #             cases2[case_key].append(case_value)
    #             cases[mach_number][weight_config][store_config].append(case_value)
    #             continue
    #
    #         if 'opgwg' not in data_dict:
    #             matrices = data_dict['matrices']
    #             log.warning(f'data_dict_keys={list(data_dict)}; matrices_keys={list(matrices)}')
    #             # asdf
    #             mass = np.full(1, np.nan)
    #             cg = np.full(3, np.nan)
    #             inertia = np.full((3, 3), np.nan)
    #         else:
    #             opgwg = data_dict['opgwg']  # grid point weight
    #             # matrices = data_dict['matrices']
    #             # frequencies = matrices['freq']
    #             mass = opgwg['mass']
    #             cg = opgwg['cg']
    #             # print(opgwg)
    #             # print(f'frequencies = {frequencies.round(3)}')
    #             inertia = opgwg['I(S)']
    #
    #         resp = resp_dict[1]
    #         modes = resp.modes[6:]
    #         assert modes[0] == 7, modes
    #
    #         # xcutoff doesn't apply for first 6 modes
    #         resp.nrigid_body_modes = 6
    #         resp.x_cutoff = x_cutoff
    #         resp.set_plot_settings(
    #             figsize=figsize,
    #             xtick_major_locator_multiple=[50., 50.],
    #             ytick_major_locator_multiple=[0.02, None],
    #         )
    #         resp.set_symbol_settings(
    #             nopoints=False, show_mode_number=False, point_spacing=8,
    #             markersize=5,
    #         )
    #
    #         # Crossing = tuple[float, float, float]
    #         # if parse_hump_modes:
    #         #     vl_vf_crossing_dict, hump_vd_crossing_dict = resp.get_hump_flutter_crossings(
    #         #         damping_crossings=damping_crossings, modes=modes,
    #         #         eas_range=eas_range)
    #         #     print(hump_vd_crossing_dict)
    #         #     vl_vf_crossing_dict = hump_vd_crossing_dict
    #         # else:
    #         vl_vf_crossing_dict, vd_crossing_dict = resp.get_flutter_crossings(
    #             damping_crossings=damping_crossings2, modes=modes,
    #             eas_range=eas_range)
    #         # print(vl_vf_crossing_dict)
    #
    #         if make_pngs:
    #             fig, (damp_axes, freq_axes) = resp.plot_vg_vf(
    #                 plot_type='eas', modes=modes,
    #                 clear=False, close=False, legend=True,
    #                 xlim=eas_lim, ylim_damping=ylim_damping, ylim_freq=ylim_freq,
    #                 # ivelocity: Optional[int]=None,
    #                 v_lines=v_lines,
    #                 damping_required=0.0,
    #                 damping_limit=0.03,
    #                 ncol=ncol, freq_tol=freq_tol, freq_tol_remove=freq_tol_remove,
    #                 # plot_freq_tol_filtered_lines=True,
    #                 damping_crossings=damping_crossings2, filter_damping=True,
    #                 eas_range=eas_range,
    #                 png_filename=None,
    #                 filter_freq=True,
    #                 show_detailed_mode_info=show_detailed_mode_info,
    #                 show=False)
    #             basename = os.path.basename(png_filename)
    #             damp_axes.set_title(basename)
    #             # plt.tight_layout()
    #             # if show_individual:
    #             #     plt.show()
    #             fig.savefig(png_filename, bbox_inches='tight')
    #             # bbox_to_anchor=(1, 1), borderaxespad=0)
    #             shutil.copyfile(png_filename, png_filename_mach)
    #             plt.close()
    #         if show_individual:
    #             raise RuntimeError('stopping')
    #
    #         if show_detailed_mode_info and 0:
    #             keys = list(filename_base_to_vl_vf_vbase_dict)
    #             # print('filename_base_to_vl_vf_vbase_dict.keys = ', keys)
    #             # print(f'f06_filename_base = {f06_filename_base}')
    #             if len(filename_base_to_vl_vf_vbase_dict):
    #                 v0i, vfi, vbase = filename_base_to_vl_vf_vbase_dict[f06_filename_base]
    #                 print(f'v0i={v0i}, vfi={vfi}, vbase={vbase}')
    #                 vl_array, vf_array = resp.hump_modes_from_VL_VF_dict(
    #                     vl_vf_crossing_dict, v0i, vfi, vbase, log)
    #                 if len(vl_array):
    #                     print('vl:')
    #                     print(vl_array)
    #                 if len(vf_array):
    #                     print('vf:')
    #                     print(vf_array)
    #
    #         # print(f'modes = {modes}')
    #         # print(f'xcrossing_dict = {xcrossing_dict}')
    #         hump_message = _get_hump_message(
    #             resp, vl_vf_crossing_dict,
    #             eas_range, show_detailed_mode_info)
    #
    #         log.info(f'VL_target = {VL_target}')
    #         VL, freqL, VF, freqF, VD, freqD = resp.xcrossing_dict_to_VL_VF_VD(
    #             vl_vf_crossing_dict, vd_crossing_dict,
    #             log, freq_target, VL_target, VF_target,
    #             v_baseline=V_baseline,
    #             # is_hump_modes=parse_hump_modes,
    #         )
    #         if VL < VL_target:
    #             log.error(f'VL={VL} KEAS, freq={freqL} Hz; {f06_filename_base}')
    #         if VF < VF_target:
    #             log.error(f'VF={VF} KEAS, freq={freqF} Hz; {f06_filename_base}')
    #         if VD < VD_target:
    #             log.error(f'VD={VD} KEAS, freq={freqD} Hz; {f06_filename_base}')
    #
    #         case_key = (mach_number, weight_config)
    #         case_value = (VL, freqL, VF, freqF, VD,
    #                       mass, cg, inertia, configi,
    #                       hump_message,
    #                       f06_filename_base, f06_filename, png_filename)
    #         cases2[case_key].append(case_value)
    #         cases[mach_number][weight_config].append(case_value)
    #         log.debug(f'mach={mach_number} VL={VL:g} VF={VF:g} VD={VD:g} {f06_filename_base}')
    #         if hump_message:
    #             log.warning(hump_message)
    #             # asdf
    #         # print(xcrossing_dict)
    #         # return cases
    #
    # names = ['Mach', 'Weight', 'Config']
    # case_data = CaseData(cases, cases2, names)
    # return case_data

def _get_hump_message(resp: FlutterResponse,
                      vl_vf_crossing_dict,
                      eas_range: tuple[float, float],
                      show_detailed_mode_info: bool) -> str:
    hump_message = ''
    if not show_detailed_mode_info:
        return hump_message

    hump_message_list = resp.get_hump_mode_messages(
        vl_vf_crossing_dict,
        modes=None,
        eas_range=eas_range,
        filter_damping=False,
        plot_type='eas',
    )
    if hump_message_list:
        hump_message_list2 = [line.strip().replace('\n', '; ') for line in hump_message_list]
        hump_message = '\n'.join(hump_message_list2)
        # log.warning(f'hump_message = {hump_message!r}')
    return hump_message


def _write_name_value_table(document, records: dict[str, float]) -> None:
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Value'
    for name, value in records.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(name)
        row_cells[1].text = str(value)


def _write_2d_table(document: Document,
                    records: dict[str, list[list[float]]],
                    log: SimpleLogger, table_name: str) -> None:
    """
    Expects a 2D table (e.g., list of lists)

    Parameters
    ----------
    document : Document
        the word doc
    records: dict[str, list[float]
        the data to add

    Returns
    -------
    mass_table = {
        'Configuration': ['Clean'],
        'Mass (in)': [1.0],
        'XCG (in)': [1.0],
        'YCG (in)': [1.0],
        'ZCG (in)': [1.0],
        'Ixx (slinch-in^2)': [1.0],
        'Iyy (slinch-in^2)': [1.0],
        'Izz (slinch-in^2)': [1.0],
    }
    """
    keys = list(records.keys())
    nkeys = len(keys)
    assert nkeys > 0, keys
    key0 = keys[0]
    column0 = records[key0]
    ncols = nkeys
    nrows = len(column0) + 1

    table = document.add_table(rows=nrows, cols=ncols)

    # write the column headers
    hdr_cells = table.rows[0].cells
    for ikey, key in enumerate(keys):
        hdr_cells[ikey].text = key

    # write the rows
    icol = 0
    for key, values in records.items():
        for irow in range(1, nrows):
            # print(key, irow, nrows, values)
            try:
                value = values[irow-1]
            except IndexError:
                log.error(f'problem writing {table_name} for irow={irow}')
                continue
            row_cells = table.rows[irow].cells
            row_cells[icol].text = str(value)
        icol += 1
    return


def str_limit_to_limit(data: list[str]) -> Limit:
    data_out = []
    for value in data:
        value = value.strip()
        if len(value) == 0:
            data_out.append(None)
        else:
            data_out.append(float(value))
    return data_out

def double_or_blank(value: float | str,
                    default: float=0.0) -> float:
    if isinstance(value, str):
        value = value.strip()
        if len(value) == 0:
            return default
        value2 = float(value)
        return value2
    return float(value)


def remove_empty_rows(configs: list[str],
                      f06_filenames: list[str],
                      log: SimpleLogger) -> tuple[list[str], list[str]]:
    configs2 = []
    f06_filenames2 = []
    for config, f06_filename in zip(configs, f06_filenames):
        if not os.path.exists(f06_filename):
            log.warning(print_bad_path(f06_filename))
            continue
        configs2.append(config)
        f06_filenames2.append(f06_filename)
    configs = configs2
    f06_filenames = f06_filenames2
    return configs, f06_filenames


if __name__ == '__main__':  # pragma: no cover
    # out = split_by_pattern(['cat_dog_1.0_asdf', 'cat_dog_2.0_qwer'])
    # print(out)
    main()
